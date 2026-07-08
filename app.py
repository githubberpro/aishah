import os
import re
import json
import traceback
import io
import hashlib
import hmac
from datetime import date, datetime, timedelta

try:
    import pdfplumber
    _HAS_PDF = True
except ImportError:
    _HAS_PDF = False

try:
    from docx import Document as _DocxDocument
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False

try:
    import openpyxl as _openpyxl
    _HAS_XLSX = True
except ImportError:
    _HAS_XLSX = False

try:
    import xlrd as _xlrd
    _HAS_XLS = True
except ImportError:
    _HAS_XLS = False

try:
    import streamlit as _st_pre
    for _key in ("DATABASE_URL", "ANTHROPIC_API_KEY"):
        if _key in _st_pre.secrets:
            os.environ.setdefault(_key, _st_pre.secrets[_key])
except Exception:
    pass

import streamlit as st
import pandas as pd
import plotly.graph_objects as go
import anthropic

import database as db

# ── Page config ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Aishah · AI Advisor CRM",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── Password gate ─────────────────────────────────────────────────────────────
def _auth_token(pwd: str) -> str:
    """Stable HMAC token derived from the password — stored in URL to survive reconnects."""
    return hmac.new(pwd.encode(), b"aishah-crm-v1", hashlib.sha256).hexdigest()[:40]

def _check_password() -> bool:
    _pwd_secret = st.secrets.get("APP_PASSWORD", "")
    if not _pwd_secret:
        return True  # dev mode — no password set

    _token = _auth_token(_pwd_secret)

    # Fast path: already authenticated this session
    if st.session_state.get("_authenticated"):
        return True

    # Persistent path: valid token in URL (survives WebSocket reconnects & tab sleeps)
    if st.query_params.get("auth") == _token:
        st.session_state["_authenticated"] = True
        return True

    # Show login form
    st.markdown("""
    <div style='max-width:380px;margin:80px auto 0;text-align:center'>
        <div style='font-size:3rem;margin-bottom:8px'>🤖</div>
        <h2 style='color:#33475B;margin-bottom:4px'>AISHAH</h2>
        <p style='color:#7C98B6;margin-bottom:28px;font-size:14px'>AI Smart Hustles and Alliances Hub</p>
    </div>
    """, unsafe_allow_html=True)

    _col = st.columns([1, 2, 1])[1]
    with _col:
        _pw = st.text_input("Password", type="password", placeholder="Enter password",
                            label_visibility="collapsed")
        if st.button("Sign in", type="primary", use_container_width=True):
            if _pw == _pwd_secret:
                st.session_state["_authenticated"] = True
                st.query_params["auth"] = _token  # Embed token in URL for future reconnects
                st.rerun()
            else:
                st.error("Incorrect password — please try again.")
    st.stop()
    return False

_check_password()

# ── One-time DB initialisation (cached per server process, not per page load) ──
@st.cache_resource
def _init_db_once():
    db.init_db()
    db.delete_demo_data()
    return True

_init_db_once()

# ── HubSpot-inspired CSS ──────────────────────────────────────────────────────
st.markdown("""
<style>
/* Sidebar */
[data-testid="stSidebar"] { background: #2D3E50 !important; }
[data-testid="stSidebar"] p,
[data-testid="stSidebar"] span,
[data-testid="stSidebar"] label,
[data-testid="stSidebar"] div { color: #B8D4E3 !important; }
[data-testid="stSidebar"] h1,
[data-testid="stSidebar"] h2,
[data-testid="stSidebar"] h3 { color: #FFFFFF !important; }
[data-testid="stSidebar"] hr { border-color: rgba(255,255,255,0.12) !important; }
[data-testid="stSidebar"] .stRadio label { color: #B8D4E3 !important; }
[data-testid="stSidebar"] [data-testid="stMetricValue"] { color: #FFFFFF !important; font-size: 18px !important; }
[data-testid="stSidebar"] [data-testid="stMetricLabel"] { color: #B8D4E3 !important; }

/* Metric cards */
[data-testid="metric-container"] {
    background: white;
    border-radius: 8px;
    padding: 16px 20px;
    box-shadow: 0 1px 3px rgba(0,0,0,0.08);
    border: 1px solid #E5EDF5;
}

/* Tabs */
.stTabs [aria-selected="true"] {
    color: #FF7A59 !important;
    border-bottom-color: #FF7A59 !important;
    font-weight: 600 !important;
}

/* Expander */
.streamlit-expanderHeader {
    background: white !important;
    border: 1px solid #DFE3EB !important;
    border-radius: 6px !important;
    font-weight: 600 !important;
    color: #33475B !important;
}

/* Inputs */
.stTextInput > div > div > input,
.stNumberInput > div > div > input,
.stTextArea > div > div > textarea,
.stSelectbox > div > div > div {
    border-color: #CBD6E2 !important;
    border-radius: 4px !important;
    color: #33475B !important;
    background: white !important;
}

/* HubSpot deal card */
.hs-deal {
    background: white;
    border: 1px solid #DFE3EB;
    border-radius: 6px;
    padding: 12px 14px;
    margin-bottom: 6px;
    transition: box-shadow 0.15s, border-color 0.15s;
}
.hs-deal:hover { box-shadow: 0 3px 10px rgba(0,0,0,0.10); border-color: #B0C4D8; }

/* Kanban column */
.hs-kanban-col {
    background: #F5F8FA;
    border-radius: 8px;
    padding: 10px;
    min-height: 200px;
}

/* Company card */
.hs-company {
    background: white;
    border: 1px solid #DFE3EB;
    border-radius: 8px;
    padding: 16px 20px;
    margin-bottom: 10px;
    box-shadow: 0 1px 3px rgba(0,0,0,0.06);
}

/* Section panel */
.hs-panel {
    background: white;
    border: 1px solid #DFE3EB;
    border-radius: 8px;
    padding: 20px;
    margin-bottom: 16px;
    box-shadow: 0 1px 3px rgba(0,0,0,0.06);
}

/* Badge */
.hs-badge {
    display: inline-block;
    padding: 2px 9px;
    border-radius: 10px;
    font-size: 10px;
    font-weight: 700;
    text-transform: uppercase;
    letter-spacing: 0.5px;
}

/* Won / Lost rows */
.hs-won { background:#F0FBF4; border-left:4px solid #00A862; border-radius:6px; padding:8px 12px; margin:3px 0; font-size:12px; color:#33475B; }
.hs-lost { background:#F9F9F9; border-left:4px solid #DFE3EB; border-radius:6px; padding:8px 12px; margin:3px 0; font-size:12px; color:#7C98B6; }

/* Activity timeline */
.hs-activity { border-left: 3px solid #DFE3EB; padding: 0 0 16px 16px; margin-left: 8px; position: relative; }
.hs-activity::before { content: ""; width: 10px; height: 10px; background: #FF7A59; border-radius: 50%; position: absolute; left: -7px; top: 2px; }

/* Overdue warning */
.hs-overdue { background: #FEF3F0; border: 1px solid #FBDAD2; border-radius: 6px; padding: 10px 14px; color: #C0392B; font-size: 13px; font-weight: 500; margin-bottom: 12px; }

/* Aishah FAB */
.aishah-fab {
    position: fixed; bottom: 28px; left: 28px;
    width: 58px; height: 58px; border-radius: 50%;
    background: linear-gradient(135deg, #FF7A59 0%, #FF6B45 100%);
    display: flex; align-items: center; justify-content: center;
    font-size: 26px; text-decoration: none;
    box-shadow: 0 4px 20px rgba(255,122,89,0.45);
    z-index: 9999; transition: transform 0.18s ease, box-shadow 0.18s ease;
}
.aishah-fab:hover { transform: scale(1.10); box-shadow: 0 6px 26px rgba(255,122,89,0.65); }
</style>
""", unsafe_allow_html=True)

# ── Constants ─────────────────────────────────────────────────────────────────
PAGES = ["Dashboard", "Pipeline", "Active Projects", "Account Plan", "Channel Partners", "Companies", "Activities", "Eminence"]
ICONS = {"Dashboard": "📊", "Pipeline": "💼", "Active Projects": "🚧", "Account Plan": "📋", "Channel Partners": "🤝", "Companies": "🏢", "Activities": "📅", "Eminence": "🌟"}
STAGE_ORDER = ["Prospect", "Qualified", "Proposal", "Negotiation", "Won", "Lost"]
STAGE_COLOR = {
    "Prospect": "#516F90", "Qualified": "#0091AE", "Proposal": "#F5A623",
    "Negotiation": "#E8692A", "Won": "#00A862", "Lost": "#9B9B9B",
}
STAGE_BG = {
    "Prospect": "#EAF0F6", "Qualified": "#E5F5F8", "Proposal": "#FEF5E7",
    "Negotiation": "#FEF0E8", "Won": "#E6F7ED", "Lost": "#F5F5F5",
}
SECTOR_LIST = [
    "Financial Services", "Government", "Technology", "Healthcare",
    "Logistics & Supply Chain", "Real Estate", "FMCG & Agribusiness",
    "Telecommunications", "Energy", "Education", "Other",
]
SERVICE_TYPES = [
    "Strategy & Roadmap", "GenAI Implementation", "AI Governance",
    "AI Implementation", "Workforce & Change", "Data & Analytics",
    "AI Audit & Assurance", "Other",
]
_CHAT_MODELS = {
    "Sonnet 4.6": "claude-sonnet-4-6",
    "Haiku 4.5 (Fast)": "claude-haiku-4-5-20251001",
    "Opus 4.8 (Powerful)": "claude-opus-4-8",
}
_CHAT_MODEL_DEFAULT = "Sonnet 4.6"

CONTRACT_TYPES = [
    "Project", "Retainer", "Project + Retainer",
    "Government Project", "Government Advisory", "Multi-year Program",
]
ACTIVITY_TYPES = [
    "Meeting", "Call", "Email", "Workshop", "Presentation", "Demo",
    "Proposal Submission", "Contract Signing", "Site Visit", "Event", "Note",
]
ACTIVITY_ICONS = {
    "Meeting": "🤝", "Call": "📞", "Email": "📧", "Workshop": "🛠️",
    "Presentation": "📊", "Demo": "💻", "Proposal Submission": "📋",
    "Contract Signing": "✍️", "Site Visit": "🏢", "Event": "🎪", "Note": "📝",
}

_SSIC_TO_SECTOR = {
    "641": "Financial Services", "642": "Financial Services", "643": "Financial Services",
    "649": "Financial Services", "651": "Financial Services", "652": "Financial Services",
    "661": "Financial Services", "662": "Financial Services", "663": "Financial Services",
    "620": "Technology", "621": "Technology", "630": "Technology", "631": "Technology",
    "841": "Government", "842": "Government", "843": "Government",
    "861": "Healthcare", "862": "Healthcare", "863": "Healthcare", "869": "Healthcare",
    "871": "Healthcare", "872": "Healthcare", "879": "Healthcare",
    "681": "Real Estate", "682": "Real Estate", "683": "Real Estate",
    "411": "Real Estate", "412": "Real Estate", "413": "Real Estate",
    "491": "Logistics & Supply Chain", "492": "Logistics & Supply Chain",
    "501": "Logistics & Supply Chain", "502": "Logistics & Supply Chain",
    "511": "Logistics & Supply Chain", "521": "Logistics & Supply Chain",
    "522": "Logistics & Supply Chain", "529": "Logistics & Supply Chain",
    "531": "Logistics & Supply Chain", "532": "Logistics & Supply Chain",
    "611": "Telecommunications", "612": "Telecommunications", "613": "Telecommunications",
    "619": "Telecommunications",
    "461": "FMCG & Agribusiness", "462": "FMCG & Agribusiness",
    "471": "FMCG & Agribusiness", "472": "FMCG & Agribusiness",
    "011": "FMCG & Agribusiness", "012": "FMCG & Agribusiness",
    "351": "Energy", "352": "Energy", "353": "Energy", "360": "Energy",
    "851": "Education", "852": "Education", "853": "Education", "854": "Education",
    "855": "Education", "856": "Education",
}

# ── Helpers ───────────────────────────────────────────────────────────────────

def _fmt_sgd(val):
    val = float(val or 0)
    if val >= 1_000_000:
        return f"S${val/1e6:.1f}M"
    if val >= 1_000:
        return f"S${val/1e3:.0f}K"
    return f"S${val:,.0f}"


def _days_until(date_str):
    if not date_str:
        return None
    try:
        return (datetime.strptime(date_str, "%Y-%m-%d").date() - date.today()).days
    except ValueError:
        return None


def _close_badge(date_str):
    d = _days_until(date_str)
    if d is None:
        return ""
    if d < 0:
        return f"🔴 {abs(d)}d overdue"
    if d <= 14:
        return f"🟠 {d}d"
    if d <= 30:
        return f"🟡 {d}d"
    return f"🟢 {date_str}"


def _action_badge(date_str):
    d = _days_until(date_str)
    if d is None:
        return ""
    if d < 0:
        return "⚠️ overdue"
    if d <= 3:
        return f"🔴 {d}d"
    if d <= 7:
        return f"🟡 {d}d"
    return f"🟢 {d}d"


def _stage_badge_html(stage):
    c = STAGE_COLOR.get(stage, "#888")
    bg = STAGE_BG.get(stage, "#F5F5F5")
    return f"<span class='hs-badge' style='background:{bg};color:{c}'>{stage}</span>"


def _ssic_to_sector(ssic: str) -> str:
    for prefix in (ssic[:3], ssic[:2], ssic[:1]):
        if prefix in _SSIC_TO_SECTOR:
            return _SSIC_TO_SECTOR[prefix]
    return "Other"


@st.cache_data(ttl=300, show_spinner=False)
def _acra_search(query: str):
    if len(query.strip()) < 2:
        return []
    try:
        import requests
        ca = os.environ.get("SSL_CERT_FILE", "/root/.ccr/ca-bundle.crt")
        verify = ca if os.path.exists(ca) else True
        resp = requests.get(
            "https://data.gov.sg/api/action/datastore_search",
            params={"resource_id": "d_f5ded86ae4cdff2e16ee65dc5cc34cb", "q": query.strip(), "limit": 8},
            timeout=6, verify=verify,
        )
        if resp.ok:
            data = resp.json()
            if data.get("success"):
                return data.get("result", {}).get("records", [])
    except Exception:
        pass
    try:
        import requests
        ca = os.environ.get("SSL_CERT_FILE", "/root/.ccr/ca-bundle.crt")
        verify = ca if os.path.exists(ca) else True
        resp = requests.get(
            "https://data.gov.sg/api/action/datastore_search",
            params={"resource_id": "5a3ce3c5-d1f9-4e83-b08e-8d8e6e1e8e0e", "q": query.strip(), "limit": 8},
            timeout=6, verify=verify,
        )
        if resp.ok and resp.json().get("success"):
            return resp.json().get("result", {}).get("records", [])
    except Exception:
        pass
    return []


# ── Aishah tools ──────────────────────────────────────────────────────────────

TOOLS = [
    {"name": "get_pipeline_summary", "description": "Get the deal pipeline grouped by stage.",
     "input_schema": {"type": "object", "properties": {}, "required": []}},
    {"name": "get_clients", "description": "List companies, optionally filtered by sector or buyer type.",
     "input_schema": {"type": "object", "properties": {
         "sector": {"type": "string"},
         "buyer_type": {"type": "string", "enum": ["Institutional", "Owner"]},
     }, "required": []}},
    {"name": "get_opportunities", "description": "List deals, optionally filtered by stage.",
     "input_schema": {"type": "object", "properties": {
         "stage": {"type": "string", "enum": STAGE_ORDER},
     }, "required": []}},
    {"name": "get_upcoming_actions", "description": "Get deals with actions due in the next N days.",
     "input_schema": {"type": "object", "properties": {"days": {"type": "integer"}}, "required": []}},
    {"name": "get_engagements", "description": "Get recent activity history.",
     "input_schema": {"type": "object", "properties": {"limit": {"type": "integer"}}, "required": []}},
    {"name": "get_eminence", "description": "Get eminence items.",
     "input_schema": {"type": "object", "properties": {
         "type_filter": {"type": "string", "enum": ["Publication", "Speaking", "Event", "Award", "Media", "Advisory"]},
     }, "required": []}},
    {"name": "add_client", "description": "Add a new company to the CRM.",
     "input_schema": {"type": "object", "properties": {
         "company": {"type": "string"}, "sector": {"type": "string"},
         "sub_sector": {"type": "string"}, "company_size": {"type": "string"},
         "buyer_type": {"type": "string", "enum": ["Institutional", "Owner"]},
         "country": {"type": "string"}, "key_contact": {"type": "string"},
         "contact_title": {"type": "string"},
         "relationship_score": {"type": "integer", "minimum": 1, "maximum": 5},
         "ai_maturity": {"type": "string"}, "notes": {"type": "string"},
     }, "required": ["company", "sector", "buyer_type"]}},
    {"name": "add_opportunity", "description": "Add a new deal to the CRM.",
     "input_schema": {"type": "object", "properties": {
         "client_id": {"type": "integer"}, "title": {"type": "string"},
         "value_sgd": {"type": "number"},
         "stage": {"type": "string", "enum": STAGE_ORDER},
         "ai_service_type": {"type": "string"},
         "probability": {"type": "integer", "minimum": 0, "maximum": 100},
         "expected_close_date": {"type": "string"}, "contract_type": {"type": "string"},
         "next_action": {"type": "string"}, "next_action_date": {"type": "string"},
         "notes": {"type": "string"},
     }, "required": ["client_id", "title", "stage"]}},
    {"name": "add_engagement", "description": "Log an activity against a company or deal.",
     "input_schema": {"type": "object", "properties": {
         "client_id": {"type": "integer"}, "opportunity_id": {"type": "integer"},
         "activity_type": {"type": "string"},
         "activity_date": {"type": "string", "description": "YYYY-MM-DD"},
         "participants": {"type": "string"}, "summary": {"type": "string"},
         "outcomes": {"type": "string"}, "next_steps": {"type": "string"},
     }, "required": ["client_id", "activity_type", "activity_date"]}},
    {"name": "update_opportunity_stage", "description": "Move a deal to a new pipeline stage.",
     "input_schema": {"type": "object", "properties": {
         "opportunity_id": {"type": "integer"},
         "stage": {"type": "string", "enum": STAGE_ORDER},
         "notes": {"type": "string"},
     }, "required": ["opportunity_id", "stage"]}},
    {"name": "add_eminence", "description": "Record a publication, speaking engagement, award, or media appearance.",
     "input_schema": {"type": "object", "properties": {
         "type": {"type": "string", "enum": ["Publication", "Speaking", "Event", "Award", "Media", "Advisory"]},
         "title": {"type": "string"}, "date": {"type": "string", "description": "YYYY-MM-DD"},
         "sector": {"type": "string"}, "platform": {"type": "string"},
         "description": {"type": "string"},
         "impact_score": {"type": "integer", "minimum": 1, "maximum": 5},
         "url": {"type": "string"},
     }, "required": ["type", "title", "date"]}},
]


def _execute_tool(name: str, inputs: dict):
    try:
        if name == "get_pipeline_summary":
            return db.get_pipeline_summary()
        elif name == "get_clients":
            return db.get_clients(sector=inputs.get("sector"), buyer_type=inputs.get("buyer_type"))
        elif name == "get_opportunities":
            return db.get_opportunities(stage=inputs.get("stage"))
        elif name == "get_upcoming_actions":
            return db.get_upcoming_actions(days=inputs.get("days", 14))
        elif name == "get_engagements":
            return db.get_engagements(limit=inputs.get("limit", 10))
        elif name == "get_eminence":
            return db.get_eminence(type_filter=inputs.get("type_filter"))
        elif name == "add_client":
            safe = {k: inputs.get(k, v) for k, v in [
                ("company", ""), ("sector", ""), ("sub_sector", ""), ("company_size", ""),
                ("buyer_type", "Institutional"), ("country", "Singapore"),
                ("key_contact", ""), ("contact_title", ""),
                ("relationship_score", 3), ("ai_maturity", ""), ("notes", ""),
            ]}
            if inputs.get("id"):
                safe["id"] = inputs["id"]
            return {"success": True, "client_id": db.upsert_client(safe)}
        elif name == "add_opportunity":
            safe = {k: inputs.get(k, v) for k, v in [
                ("client_id", None), ("title", ""), ("value_sgd", 0),
                ("stage", "Prospect"), ("ai_service_type", ""), ("probability", 20),
                ("expected_close_date", ""), ("contract_type", ""),
                ("next_action", ""), ("next_action_date", ""), ("notes", ""),
            ]}
            if inputs.get("id"):
                safe["id"] = inputs["id"]
            return {"success": True, "opportunity_id": db.upsert_opportunity(safe)}
        elif name == "add_engagement":
            safe = {k: inputs.get(k, v) for k, v in [
                ("client_id", None), ("opportunity_id", None), ("activity_type", ""),
                ("activity_date", ""), ("participants", ""), ("summary", ""),
                ("outcomes", ""), ("next_steps", ""),
            ]}
            return {"success": True, "engagement_id": db.add_engagement(safe)}
        elif name == "update_opportunity_stage":
            opp = db.get_opportunity(inputs["opportunity_id"])
            if opp:
                opp["stage"] = inputs["stage"]
                if inputs.get("notes"):
                    opp["notes"] = (opp.get("notes") or "") + f"\n[{date.today()}] {inputs['notes']}"
                db.upsert_opportunity(opp)
                return {"success": True}
            return {"error": "Opportunity not found"}
        elif name == "add_eminence":
            safe = {k: inputs.get(k, v) for k, v in [
                ("type", ""), ("title", ""), ("date", ""), ("sector", ""),
                ("platform", ""), ("description", ""), ("impact_score", 3), ("url", ""),
            ]}
            return {"success": True, "eminence_id": db.add_eminence(safe)}
        return {"error": f"Unknown tool: {name}"}
    except Exception as e:
        return {"error": f"{name} failed: {e}", "traceback": traceback.format_exc()}


_AISHAH_SYSTEM = (
    "You are Aishah — AI for Smart Hustle and Alliances Hub. You are the personal AI strategist "
    "for a senior PwC Singapore partner building toward being Singapore's #1 AI advisor.\n\n"
    "Your capabilities:\n"
    "- Query and update the CRM: companies, deals/pipeline, activities, eminence\n"
    "- Give strategic advice on winning deals, building relationships, and growing eminence\n"
    "- Analyse pipeline health, forecast revenue, surface at-risk deals\n"
    "- Help craft outreach strategies, proposals, and negotiation approaches\n\n"
    "Key context:\n"
    "- INSTITUTIONAL buyers = C-suite at large listed corps (budget buyers, institutional trust)\n"
    "- OWNER buyers = entrepreneurs, founders, family offices (wealth buyers, personal trust)\n"
    "- Owner relationships are portable post-PwC — prioritise them\n"
    "- Pipeline: Prospect → Qualified → Proposal → Negotiation → Won / Lost\n"
    "- Singapore focus; target senior decision-makers at large corporates, govt agencies, and owner-led businesses\n\n"
    "Always call tools to get real CRM data before answering. Be strategic, concise, actionable. "
    f"Today is {date.today().isoformat()}."
)

_EMINENCE_SYSTEM = (
    "You are Aishah, the eminence coach for a senior PwC Singapore AI advisory partner. "
    "Your sole focus here is helping build their public profile as Singapore's #1 practical AI advisor.\n\n"
    "Their positioning: 'Grounded Strategist' — someone whose strategic advice is forged from real deployment experience, "
    "not frameworks. 6 live AI agent deployments (Nippon Paint, SETSCO, IMC, NanoFilm, Everllence, Megastar), "
    "ECI grant approved across multiple clients, proprietary 'AI in a Box' framework, "
    "two sector beachheads: Industrial/Manufacturing and Maritime.\n\n"
    "Four content archetypes to rotate:\n"
    "1. Field Notes — specific deployment story, show the mess, end with the lesson\n"
    "2. Pattern — synthesise across deployments, what nobody else sees\n"
    "3. Board Implication — CEO/board-level strategic takeaway\n"
    "4. Contrarian Bet — the thing you believed before the crowd, and what happened\n\n"
    "Eminence 100/200/300-day plan:\n"
    "100-day: Mine deployment gold (AI-in-a-Box framework doc, 2 client outcome stories), pick public lane "
    "(industrial/maritime AI), ship 1 LinkedIn essay/week, warm SBF/NTU/SLA ecosystem.\n"
    "200-day: 2-3 speaking slots (SBF, industry associations), co-create NTU/SBF AI-adoption benchmark study, "
    "first BT/ST op-ed, build maritime beachhead with IMC/Megastar.\n"
    "300-day: Annual 'State of Practical AI in Singapore's Industrial Economy' report, "
    "host invite-only 15-20 leader roundtable, regionalise one Singapore story to Asia stage.\n\n"
    "PR rules: Evidence over opinion. Generosity over self-reference. Consistency over intensity. "
    "4:1 ratio — 4 posts about client/sector/reader, 1 about you. "
    "Never say 'I am a top AI advisor' — make clients' results and readers look good until others say it.\n\n"
    "Reference anchors: Andrew Ng (practitioner who built category), Azeem Azhar (newsletter-to-authority), "
    "Kai-Fu Lee (Asia AI narrative owner), cite MIT Sloan/Stanford HAI/MGI data then localise.\n\n"
    "Be direct, specific, and actionable. Draft actual LinkedIn posts, op-ed openings, and speaking abstracts "
    "when asked. Pull from their real deal data when relevant. "
    f"Today is {date.today().isoformat()}."
)


# ── Aishah dialog ─────────────────────────────────────────────────────────────

@st.dialog("🤖 Aishah · AI Advisor", width="large")
def _aishah_dialog():
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        st.error("Add `ANTHROPIC_API_KEY` to your Streamlit secrets to activate Aishah.")
        return

    client_ai = anthropic.Anthropic(api_key=api_key)

    history = db.get_chat_history(limit=40)
    with st.container(height=400):
        for msg in history:
            with st.chat_message(msg["role"], avatar="🤖" if msg["role"] == "assistant" else None):
                st.markdown(msg["content"])

    _clr_col, _mdl_col = st.columns([1, 3])
    with _clr_col:
        if st.button("🗑️ Clear", help="Clear chat history", use_container_width=True):
            db.clear_chat_history()
            st.rerun()
    _aishah_model = _CHAT_MODELS[_mdl_col.selectbox(
        "Model", list(_CHAT_MODELS.keys()),
        index=list(_CHAT_MODELS.keys()).index(_CHAT_MODEL_DEFAULT),
        key="aishah_model_sel", label_visibility="collapsed",
    )]

    prompt = st.chat_input("Ask Aishah — strategy, pipeline, who to call next...")

    if prompt:
        with st.chat_message("user"):
            st.markdown(prompt)
        db.save_message("user", prompt)

        recent = db.get_chat_history(limit=20)
        api_messages = [{"role": m["role"], "content": m["content"]} for m in recent[:-1]]
        api_messages.append({"role": "user", "content": prompt})

        _api_error = None
        final_text = ""
        tool_call_count = 0

        with st.chat_message("assistant", avatar="🤖"):
            with st.spinner("Thinking..."):
                try:
                    iteration = 0
                    while iteration < 10:
                        iteration += 1
                        response = client_ai.messages.create(
                            model=_aishah_model,
                            max_tokens=8192,
                            system=_AISHAH_SYSTEM,
                            tools=TOOLS,
                            messages=api_messages,
                        )
                        assistant_content: list = []
                        text_parts: list = []
                        tool_uses: list = []
                        for block in response.content:
                            if hasattr(block, "type") and block.type == "text":
                                text_parts.append(block.text)
                                assistant_content.append({"type": "text", "text": block.text})
                            elif hasattr(block, "type") and block.type == "tool_use":
                                tool_uses.append(block)
                                assistant_content.append({
                                    "type": "tool_use", "id": block.id,
                                    "name": block.name, "input": block.input,
                                })
                        api_messages.append({"role": "assistant", "content": assistant_content})
                        if response.stop_reason in ("end_turn", "max_tokens") or not tool_uses:
                            final_text = "\n".join(text_parts)
                            break
                        tool_results = []
                        for tu in tool_uses:
                            result = _execute_tool(tu.name, tu.input)
                            tool_call_count += 1
                            tool_results.append({
                                "type": "tool_result", "tool_use_id": tu.id,
                                "content": json.dumps(result, default=str),
                            })
                        api_messages.append({"role": "user", "content": tool_results})
                except anthropic.APIStatusError as exc:
                    _api_error = f"API error {exc.status_code}: {exc.message}"
                except anthropic.APIConnectionError:
                    _api_error = "Could not reach the Claude API — check network or try again."
                except Exception as exc:
                    _api_error = f"Unexpected error: {exc}\n\n```\n{traceback.format_exc()}\n```"

            if _api_error:
                final_text = f"⚠️ {_api_error}"
            elif not final_text:
                final_text = (
                    f"Done — completed {tool_call_count} CRM action(s)."
                    if tool_call_count else "I didn't get a response. Please try again."
                )
            st.markdown(final_text)

        db.save_message("assistant", final_text)


# ── Deal form (shared across pages) ──────────────────────────────────────────

def _deal_form(o: dict, all_clients: list, form_key: str = "opp_form"):
    """Render deal create/edit form."""
    _client_map = {c["company"].lower(): c for c in all_clients}
    current_company = next((c["company"] for c in all_clients if c["id"] == o.get("client_id")), "")

    with st.form(form_key):
        st.markdown("#### 💼 " + ("New Opportunity" if not o.get("id") else "Edit Opportunity"))

        # Company as free-text — looks up existing client or auto-creates
        company_name = st.text_input(
            "Company *",
            value=current_company or o.get("company", ""),
            placeholder="Type the company name",
        )

        f1, f2 = st.columns(2)
        title = f1.text_input("Opportunity Title *", value=o.get("title", ""))
        ai_svc = f2.selectbox(
            "Service Type", SERVICE_TYPES,
            index=SERVICE_TYPES.index(o["ai_service_type"]) if o.get("ai_service_type") in SERVICE_TYPES else 0,
        )

        f3, f4, f5 = st.columns(3)
        value_sgd = f3.number_input(
            "Contract Value (S$K) *", min_value=0.0, value=float(o.get("value_sgd", 0)) / 1000,
            step=50.0, format="%.0f",
        )
        stage = f4.selectbox(
            "Stage *", STAGE_ORDER,
            index=STAGE_ORDER.index(o.get("stage", "Prospect")),
        )
        probability = f5.slider("Probability %", 0, 100, o.get("probability", 20))

        f6, f7 = st.columns(2)
        contract_type = f6.selectbox(
            "Contract Type", CONTRACT_TYPES,
            index=CONTRACT_TYPES.index(o["contract_type"]) if o.get("contract_type") in CONTRACT_TYPES else 0,
        )
        _ecd_val = None
        if o.get("expected_close_date"):
            try:
                _ecd_val = datetime.strptime(o["expected_close_date"], "%Y-%m-%d").date()
            except ValueError:
                pass
        expected_close_date = f7.date_input("Expected Close Date", value=_ecd_val)

        # Engagement team
        e1, e2 = st.columns(2)
        engagement_manager = e1.text_input(
            "EM (Engagement Manager)",
            value=o.get("engagement_manager", ""),
            placeholder="e.g. John Lim",
        )
        engagement_partner = e2.text_input(
            "EP (Engagement Partner)",
            value=o.get("engagement_partner", ""),
            placeholder="e.g. Sarah Tan",
        )

        # Decision-maker and influencers
        p1, p2 = st.columns(2)
        decision_maker = p1.text_input(
            "Decision Maker",
            value=o.get("decision_maker", ""),
            placeholder="Name · Title (e.g. Jane Tan · CFO)",
        )
        influencers = p2.text_input(
            "Influencers",
            value=o.get("influencers", ""),
            placeholder="e.g. CTO, Head of Procurement",
        )

        _psd_val = None
        if o.get("project_start_date"):
            try:
                _psd_val = datetime.strptime(o["project_start_date"], "%Y-%m-%d").date()
            except ValueError:
                pass
        _ped_val = None
        if o.get("project_end_date"):
            try:
                _ped_val = datetime.strptime(o["project_end_date"], "%Y-%m-%d").date()
            except ValueError:
                pass
        _pd1, _pd2 = st.columns(2)
        project_start_date = _pd1.date_input("Project Start Date", value=_psd_val)
        project_end_date = _pd2.date_input("Project End Date", value=_ped_val)

        next_action = st.text_input(
            "Next Action", value=o.get("next_action", ""),
            placeholder="e.g. Send proposal deck to CFO",
        )
        _nad_val = None
        if o.get("next_action_date"):
            try:
                _nad_val = datetime.strptime(o["next_action_date"], "%Y-%m-%d").date()
            except ValueError:
                pass
        next_action_date = st.date_input("Next Action Date", value=_nad_val)
        notes = st.text_area("Notes", value=o.get("notes", ""), height=80)

        s1, s2 = st.columns(2)
        saved = s1.form_submit_button("💾 Save Opportunity", type="primary", use_container_width=True)
        cancelled = s2.form_submit_button("Cancel", use_container_width=True)

        if saved and title and company_name:
            # Resolve or auto-create client
            existing = _client_map.get(company_name.strip().lower())
            if existing:
                client_id = existing["id"]
            else:
                client_id = db.upsert_client(dict(
                    company=company_name.strip(), sector="Other",
                    buyer_type="Institutional", country="Singapore",
                    relationship_score=3,
                ))
            db.upsert_opportunity(dict(
                id=o.get("id"), client_id=client_id,
                title=title, value_sgd=value_sgd * 1000, stage=stage,
                ai_service_type=ai_svc, probability=probability,
                expected_close_date=str(expected_close_date) if expected_close_date else None,
                contract_type=contract_type, next_action=next_action,
                next_action_date=str(next_action_date) if next_action_date else None,
                decision_maker=decision_maker, influencers=influencers,
                engagement_manager=engagement_manager, engagement_partner=engagement_partner,
                project_start_date=str(project_start_date) if project_start_date else None,
                project_end_date=str(project_end_date) if project_end_date else None,
                notes=notes,
            ))
            st.session_state.pop("edit_opp", None)
            st.session_state.pop("open_opp_id", None)
            st.rerun()
        if cancelled:
            st.session_state.pop("edit_opp", None)
            st.session_state.pop("open_opp_id", None)
            st.rerun()


# ── Sidebar ───────────────────────────────────────────────────────────────────

with st.sidebar:
    st.markdown("## 🤖 AISHAH")
    st.caption("AI Advisor CRM · PwC Singapore")
    st.divider()
    page = st.radio(
        "Navigate", PAGES,
        format_func=lambda p: f"{ICONS[p]}  {p}",
        label_visibility="collapsed",
        key="nav_radio",
    )
    st.divider()
    _sb_opps = db.get_opportunities()
    _sb_active = [o for o in _sb_opps if o["stage"] not in ("Won", "Lost")]
    _sb_won = [o for o in _sb_opps if o["stage"] == "Won"]
    _sb_lost = [o for o in _sb_opps if o["stage"] == "Lost"]
    _sb_closed = len(_sb_won) + len(_sb_lost)
    _sb_wr = round(len(_sb_won) / _sb_closed * 100) if _sb_closed else 0
    _sb_weighted = sum((o["value_sgd"] or 0) * (o["probability"] or 0) / 100 for o in _sb_active)
    _sb_won_val = sum(o["value_sgd"] or 0 for o in _sb_won)
    st.metric("Weighted Pipeline", _fmt_sgd(_sb_weighted))
    st.metric("Won (FY)", _fmt_sgd(_sb_won_val))
    st.metric("Win Rate", f"{_sb_wr}%")
    st.caption(f"Updated {datetime.now().strftime('%d %b %Y %H:%M')}")
    st.divider()
    if st.button("💬 Ask Aishah", type="primary", use_container_width=True):
        _aishah_dialog()

# ── Floating Aishah bubble ────────────────────────────────────────────────────

st.markdown("""
<a href="?aishah=1" target="_self" class="aishah-fab" title="Chat with Aishah">🤖</a>
""", unsafe_allow_html=True)

if st.query_params.get("aishah") == "1":
    st.query_params.clear()
    _aishah_dialog()


# ── Shared sortable table helper ──────────────────────────────────────────────

def _deals_table(opps: list, key_prefix: str, mode: str = "active"):
    """Render a selectable table; tapping a row opens its inline panel.
    mode: "active" | "won" | "projects"
    """
    if not opps:
        return

    id_list = [o["id"] for o in opps]

    def _to_date(s):
        if not s:
            return None
        try:
            return datetime.strptime(str(s)[:10], "%Y-%m-%d").date()
        except ValueError:
            return None

    rows = []
    for o in opps:
        row: dict = {
            "Company": o.get("company", ""),
            "Title": o.get("title", ""),
            "EP": o.get("engagement_partner") or "",
            "EM": o.get("engagement_manager") or "",
        }
        if mode == "active":
            row["Stage"] = o.get("stage", "")
            row["Value (S$)"] = float(o.get("value_sgd") or 0)
            row["Prob %"] = int(o.get("probability") or 0)
            row["Close Date"] = _to_date(o.get("expected_close_date"))
            row["Next Action"] = o.get("next_action") or ""
        elif mode == "won":
            row["Value (S$)"] = float(o.get("value_sgd") or 0)
            row["Close Date"] = _to_date(o.get("expected_close_date"))
            row["End Date"] = _to_date(o.get("project_end_date"))
            row["Service"] = o.get("ai_service_type") or ""
        else:  # projects
            row["Start"] = _to_date(o.get("project_start_date"))
            row["End"] = _to_date(o.get("project_end_date"))
            row["Contract (S$)"] = float(o.get("value_sgd") or 0)
            row["PwC Rev (S$)"] = float(o.get("pwc_revenue") or 0)
            row["Non-PwC (S$)"] = float(o.get("non_pwc_revenue") or 0)
            row["WIP (S$)"] = float(o.get("wip") or 0)
            row["Rev to Go (S$)"] = (
                float(o.get("value_sgd") or 0)
                - float(o.get("non_pwc_revenue") or 0)
                - float(o.get("wip") or 0)
            )
        rows.append(row)

    df = pd.DataFrame(rows)
    st.caption("Tap any row to open details and edit.")
    sel_event = st.dataframe(
        df,
        use_container_width=True,
        hide_index=True,
        on_select="rerun",
        selection_mode="single-row",
        key=f"{key_prefix}_sel",
    )
    sel_rows = sel_event.selection.rows

    # Determine which opp to show in the inline panel
    _panel_opp_id = None
    if sel_rows:
        _panel_opp_id = id_list[sel_rows[0]]
        st.session_state["open_opp_id"] = _panel_opp_id
        st.session_state.pop("edit_opp", None)
    elif st.session_state.get("open_opp_id") in id_list:
        # Panel was opened by a button elsewhere (e.g. kanban) for an opp in this list
        _panel_opp_id = st.session_state["open_opp_id"]

    if _panel_opp_id:
        st.divider()
        _opp_inline_panel(_panel_opp_id)


# ─────────────────────────────────────────────────────────────────────────────
# OPPORTUNITY DETAIL (scope upload · proposal · chat)
# ─────────────────────────────────────────────────────────────────────────────

_PWC_ROLES = ["Partner", "Director", "Senior Manager", "Manager", "Senior Associate", "Associate"]
_PWC_RATES = {"Partner": 5000, "Director": 3500, "Senior Manager": 2500,
              "Manager": 1800, "Senior Associate": 1200, "Associate": 800}

_PROPOSAL_SYSTEM = (
    "You are Aishah, a senior AI advisor's personal assistant at PwC Singapore. "
    "You help draft and refine client proposals for AI advisory and implementation engagements.\n\n"
    "PwC Singapore standard daily rates (S$):\n"
    "  Partner: $5,000 | Director: $3,500 | Senior Manager: $2,500\n"
    "  Manager: $1,800 | Senior Associate: $1,200 | Associate: $800\n\n"
    "Proposal standards:\n"
    "- Activities structured in phases (Discovery, Design, Implement, Embed)\n"
    "- Each phase has clear deliverables and a resource mix\n"
    "- Pricing = sum of (role × rate × days) across all phases\n"
    "- Be specific, commercially realistic, and aligned with PwC quality standards\n"
    "- Singapore government projects typically use PSDS guidelines\n\n"
    "When generating or updating a proposal, return a JSON object with this exact structure:\n"
    "{\n"
    '  "executive_summary": "string",\n'
    '  "scope_understanding": "string",\n'
    '  "activities": [\n'
    '    {"phase": "Phase 1: Name", "description": "string",\n'
    '     "activities": ["activity 1", "activity 2"],\n'
    '     "duration_weeks": 3, "deliverables": ["deliverable 1"]}\n'
    '  ],\n'
    '  "resources": [\n'
    '    {"role": "Partner", "rate_per_day": 5000, "days": 8}\n'
    '  ],\n'
    '  "milestones": [\n'
    '    {"week": 1, "milestone": "Project Kick-off & Mobilisation"}\n'
    '  ],\n'
    '  "pricing_breakdown": [\n'
    '    {"item": "Phase 1: Discovery", "amount": 80000, "notes": ""}\n'
    '  ],\n'
    '  "assumptions": ["string"],\n'
    '  "exclusions": ["string"]\n'
    "}\n\n"
    "When adjusting a proposal based on user feedback, return ONLY the updated JSON — no prose before or after."
)

_CP_PREP_SYSTEM = (
    "You are Aishah, the personal AI advisor assistant to a senior partner at PwC Singapore. "
    "You specialise in helping prepare for channel partner meetings — government agencies, trade bodies, "
    "statutory boards, industry associations, and institutional partners.\n\n"
    "Your role is to help the partner:\n"
    "- Craft a sharp, purposeful meeting agenda\n"
    "- Identify key talking points and value propositions to reinforce\n"
    "- Anticipate partner concerns or objections and prepare responses\n"
    "- Suggest follow-up commitments or joint initiatives to propose\n"
    "- Draft meeting notes, action items, or follow-up emails post-meeting\n"
    "- Advise on relationship health and how to deepen engagement\n\n"
    "You have access to the partner profile below. Draw on it to give specific, actionable advice "
    "rather than generic talking points. Be concise and commercially sharp — this is a busy partner "
    "who needs briefing-quality output, not essays.\n\n"
    "Tone: professional, direct, thoughtful. PwC Singapore context throughout."
)


def _extract_text(uploaded_file) -> str:
    """Extract plain text from a PDF, DOCX, TXT, or Excel upload."""
    name = uploaded_file.name.lower()
    raw = uploaded_file.read()
    if name.endswith(".txt"):
        return raw.decode("utf-8", errors="ignore")
    if name.endswith(".pdf") and _HAS_PDF:
        text_parts = []
        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        return "\n\n".join(text_parts)
    if name.endswith(".docx") and _HAS_DOCX:
        doc = _DocxDocument(io.BytesIO(raw))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    if name.endswith(".xls"):
        # Legacy binary .xls format — openpyxl can't read it, so use xlrd.
        if not _HAS_XLS:
            return "(Legacy .xls extraction unavailable — xlrd not installed)"
        try:
            book = _xlrd.open_workbook(file_contents=raw)
            parts = []
            for sheet in book.sheets():
                parts.append(f"=== Sheet: {sheet.name} ===")
                for r in range(sheet.nrows):
                    cells = [str(c) if c is not None else "" for c in sheet.row_values(r)]
                    if any(c.strip() for c in cells):
                        parts.append("\t".join(cells))
            return "\n".join(parts)
        except Exception as exc:
            return f"(Excel parse error: {exc})"
    if name.endswith(".xlsx"):
        if not _HAS_XLSX:
            return "(Excel extraction unavailable — openpyxl not installed)"
        try:
            wb = _openpyxl.load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
            parts = []
            for sheet in wb.worksheets:
                parts.append(f"=== Sheet: {sheet.title} ===")
                for row in sheet.iter_rows(values_only=True):
                    cells = [str(c) if c is not None else "" for c in row]
                    if any(c.strip() for c in cells):
                        parts.append("\t".join(cells))
            wb.close()
            return "\n".join(parts)
        except Exception as exc:
            return f"(Excel parse error: {exc})"
    return raw.decode("utf-8", errors="ignore")[:50000]


def _generate_proposal(opp: dict, scope_docs: list) -> tuple:
    """Call Claude to generate a structured proposal JSON.
    Returns (dict, None) on success or (None, error_str) on failure.
    """
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return None, "ANTHROPIC_API_KEY not set in Streamlit secrets."
    scope_text = "\n\n---\n\n".join(
        f"**{d['filename']}**\n{d['extracted_text'][:8000]}"
        for d in scope_docs
    ) if scope_docs else "(No scope documents uploaded — generate from deal information only)"

    prompt = (
        f"Opportunity: {opp.get('title','')}\n"
        f"Client: {opp.get('company','')}\n"
        f"Service Type: {opp.get('ai_service_type','')}\n"
        f"Contract Value Target: S${opp.get('value_sgd',0):,.0f}\n"
        f"Expected Close: {opp.get('expected_close_date','TBD')}\n"
        f"EP: {opp.get('engagement_partner','')}, EM: {opp.get('engagement_manager','')}\n\n"
        f"SCOPE DOCUMENTS:\n{scope_text}\n\n"
        "Generate a comprehensive PwC-quality proposal JSON."
    )
    client_ai = anthropic.Anthropic(api_key=api_key)
    # Stream the request so a long generation keeps the connection active —
    # a non-streaming call can sit idle long enough for a proxy to drop it,
    # which surfaces as the request silently stopping with no result.
    # Try adaptive thinking first; fall back to a plain request if unsupported.
    for kwargs in [
        {"thinking": {"type": "adaptive"}, "max_tokens": 8000},
        {"max_tokens": 4096},
    ]:
        try:
            with client_ai.messages.stream(
                model="claude-opus-4-8",
                system=_PROPOSAL_SYSTEM,
                messages=[{"role": "user", "content": prompt}],
                **kwargs,
            ) as stream:
                resp = stream.get_final_message()
            for block in resp.content:
                if hasattr(block, "text"):
                    txt = block.text.strip()
                    # Try ```json fence first, then bare JSON object
                    fence_start = txt.find("```json")
                    if fence_start >= 0:
                        fence_end = txt.find("```", fence_start + 7)
                        if fence_end > fence_start:
                            txt = txt[fence_start + 7:fence_end].strip()
                    start = txt.find("{")
                    end = txt.rfind("}") + 1
                    if start >= 0 and end > start:
                        return json.loads(txt[start:end]), None
            return None, "Claude responded but returned no JSON. Check the system prompt."
        except anthropic.BadRequestError as exc:
            if "thinking" in str(exc).lower() or "adaptive" in str(exc).lower():
                continue  # retry without thinking
            return None, f"API error: {exc}"
        except Exception as exc:
            return None, f"{type(exc).__name__}: {exc}"
    return None, "Generation failed after retrying without extended thinking."


def _proposal_total(proposal: dict) -> float:
    return sum(
        r.get("rate_per_day", 0) * r.get("days", 0)
        for r in proposal.get("resources", [])
    )


def _dedup_chat(msgs: list) -> list:
    """Remove consecutive same-role messages, keeping the last in each run.

    Duplicate user messages accumulate when a submission times out and the user
    retries. We keep the last message in each same-role run so the one closest
    to the following assistant reply is the one shown.
    """
    result = []
    i = 0
    while i < len(msgs):
        j = i
        while j + 1 < len(msgs) and msgs[j + 1]["role"] == msgs[i]["role"]:
            j += 1
        result.append(msgs[j])
        i = j + 1
    return result


def _cp_inline_panel(cp_id: int):
    """Inline panel with Relationship and Meeting Prep tabs for a channel partner."""
    cp = db.get_channel_partner(cp_id)
    if not cp:
        st.error("Partner not found.")
        if st.button("✖ Close", key=f"cp_panel_close_err_{cp_id}"):
            st.session_state.pop("open_cp_id", None)
            st.rerun()
        return

    _tier_colors = {"Strategic": "#C0392B", "Gold": "#D4A017", "Silver": "#7F8C8D", "Standard": "#2980B9"}
    _tier_c = _tier_colors.get(cp.get("tier", "Standard"), "#2980B9")

    with st.container(border=True):
        _cph1, _cph2, _cph3, _cph4 = st.columns([7, 1, 1, 1])
        _cph1.markdown(
            f"<h4 style='margin:0;color:#33475B'>{cp.get('name','')}</h4>"
            f"<span style='color:{_tier_c};font-weight:700;font-size:13px'>{cp.get('tier','Standard')}</span>"
            f"&nbsp;&nbsp;<span style='color:#7C98B6;font-size:13px'>"
            f"{cp.get('partner_type','')} &nbsp;·&nbsp; {cp.get('status','')} &nbsp;·&nbsp; "
            f"EP: {cp.get('engagement_partner','—')}</span>",
            unsafe_allow_html=True,
        )
        if _cph2.button("✏️ Edit", key=f"cp_edit_btn_{cp_id}", use_container_width=True):
            st.session_state["edit_cp"] = db.get_channel_partner(cp_id)
            st.rerun()
        if _cph3.button("🗑️ Del", key=f"cp_del_btn_{cp_id}", use_container_width=True,
                        help="Delete this partner"):
            db.delete_channel_partner(cp_id)
            st.session_state.pop("open_cp_id", None)
            st.rerun()
        if _cph4.button("✖ Close", key=f"cp_panel_close_{cp_id}", use_container_width=True):
            st.session_state.pop("open_cp_id", None)
            st.rerun()

        tab_rel, tab_prep = st.tabs(["📋 Relationship", "🤝 Meeting Prep"])

        # ══ TAB 1 — RELATIONSHIP ══════════════════════════════════════════════
        with tab_rel:
            st.markdown("##### Relationship Health & Meeting Log")

            _rl1, _rl2 = st.columns([1, 3])
            _health_labels = {1: "🔴 Critical", 2: "🟠 At Risk", 3: "🟡 Steady", 4: "🟢 Strong", 5: "🌟 Exemplary"}
            _cur_health = int(cp.get("health_score") or 3)
            _new_health = _rl1.selectbox(
                "Relationship Health",
                options=[1, 2, 3, 4, 5],
                index=_cur_health - 1,
                format_func=lambda x: _health_labels[x],
                key=f"cp_health_{cp_id}",
            )

            def _cp_parse_date(val):
                if val:
                    try:
                        return datetime.strptime(str(val)[:10], "%Y-%m-%d").date()
                    except ValueError:
                        pass
                return None

            _rl3, _rl4 = st.columns(2)
            _new_last = _rl3.date_input(
                "Last Meeting Date",
                value=_cp_parse_date(cp.get("last_meeting_date")),
                key=f"cp_last_mtg_{cp_id}",
            )
            _new_next = _rl4.date_input(
                "Next Meeting Date",
                value=_cp_parse_date(cp.get("next_meeting_date")),
                key=f"cp_next_mtg_{cp_id}",
            )
            _new_purpose = st.text_input(
                "Meeting Purpose / Agenda Theme",
                value=cp.get("meeting_purpose") or "",
                placeholder="e.g. Q3 pipeline review, MOU renewal discussion, joint GTM planning",
                key=f"cp_purpose_{cp_id}",
            )
            _new_notes = st.text_area(
                "Relationship Notes",
                value=cp.get("notes") or "",
                height=120,
                placeholder="Key relationship context, recent wins, sensitivities, open commitments…",
                key=f"cp_notes_{cp_id}",
            )

            if st.button("💾 Save Relationship Info", type="primary", key=f"cp_rel_save_{cp_id}"):
                db.upsert_channel_partner({
                    **cp,
                    "health_score": _new_health,
                    "last_meeting_date": str(_new_last) if _new_last else "",
                    "next_meeting_date": str(_new_next) if _new_next else "",
                    "meeting_purpose": _new_purpose,
                    "notes": _new_notes,
                })
                st.success("Saved.")
                st.rerun()

            # ── Relationship snapshot ─────────────────────────────────────────
            st.divider()
            _snap1, _snap2, _snap3 = st.columns(3)
            _snap1.metric("Referrals Received", int(cp.get("referrals_received") or 0))
            _conv = int(cp.get("referrals_converted") or 0)
            _recv = int(cp.get("referrals_received") or 0)
            _snap2.metric("Referrals Converted", _conv,
                          delta=f"{round(_conv/_recv*100)}% rate" if _recv else None)
            _snap3.metric("Joint Pipeline", _fmt_sgd(float(cp.get("joint_pipeline_value") or 0)))

        # ══ TAB 2 — MEETING PREP ══════════════════════════════════════════════
        with tab_prep:
            api_key = os.environ.get("ANTHROPIC_API_KEY", "")
            if not api_key:
                st.warning("Add `ANTHROPIC_API_KEY` to Streamlit secrets to use Meeting Prep.")
            else:
                # Build partner context for Aishah
                _next_mtg_str = cp.get("next_meeting_date") or "not scheduled"
                _last_mtg_str = cp.get("last_meeting_date") or "no record"
                _health_str = _health_labels.get(int(cp.get("health_score") or 3), "Steady")
                _cp_ctx = (
                    f"Partner: {cp.get('name','')} | Type: {cp.get('partner_type','')} | "
                    f"Tier: {cp.get('tier','')} | Status: {cp.get('status','')}\n"
                    f"Program: {cp.get('program_name','—')}\n"
                    f"Focus Sectors: {cp.get('focus_sectors','—')}\n"
                    f"Primary Contact: {cp.get('primary_contact','—')} ({cp.get('contact_title','')}) "
                    f"<{cp.get('contact_email','')}>\n"
                    f"EP: {cp.get('engagement_partner','—')} | EM: {cp.get('engagement_manager','—')}\n"
                    f"MOU Date: {cp.get('mou_date','—')} | Renewal: {cp.get('renewal_date','—')}\n"
                    f"Relationship Health: {_health_str}\n"
                    f"Last Meeting: {_last_mtg_str} | Next Meeting: {_next_mtg_str}\n"
                    f"Meeting Purpose: {cp.get('meeting_purpose','—')}\n"
                    f"Referrals: {cp.get('referrals_received',0)} received, "
                    f"{cp.get('referrals_converted',0)} converted\n"
                    f"Joint Pipeline: S${float(cp.get('joint_pipeline_value') or 0):,.0f}\n"
                    f"Notes: {cp.get('notes','—')}"
                )
                _cp_chat_system = (
                    f"{_CP_PREP_SYSTEM}\n\n"
                    f"=== PARTNER PROFILE ===\n{_cp_ctx}"
                )

                cp_history = _dedup_chat(db.get_cp_chat(cp_id, limit=40))
                _clr2, _cp_mdl_col, _ = st.columns([1, 2, 4])
                if _clr2.button("🗑️ Clear chat", key=f"clr_cp_chat_{cp_id}"):
                    db.clear_cp_chat(cp_id)
                    st.rerun()
                _cp_chat_model = _CHAT_MODELS[_cp_mdl_col.selectbox(
                    "Model", list(_CHAT_MODELS.keys()),
                    index=list(_CHAT_MODELS.keys()).index(_CHAT_MODEL_DEFAULT),
                    key=f"cp_chat_model_{cp_id}", label_visibility="collapsed",
                )]

                # Quick-start prompts
                if not cp_history:
                    st.caption("Try asking:")
                    _qs_cols = st.columns(3)
                    _quick_starts = [
                        "Draft a meeting agenda for our next session",
                        "What are the key talking points I should cover?",
                        "What follow-up actions are typically expected after this kind of meeting?",
                    ]
                    for _qi, _qs in enumerate(_quick_starts):
                        if _qs_cols[_qi].button(_qs, key=f"cp_qs_{cp_id}_{_qi}", use_container_width=True):
                            st.session_state[f"cp_prefill_{cp_id}"] = _qs
                            st.rerun()

                _prefilled = st.session_state.pop(f"cp_prefill_{cp_id}", None)
                cp_prompt = st.chat_input(
                    "Ask Aishah — meeting agenda, talking points, follow-up email, relationship advice…",
                    key=f"cp_chat_input_{cp_id}",
                )
                cp_prompt = cp_prompt or _prefilled

                _cp_chat_container = st.container(height=360)
                with _cp_chat_container:
                    for msg in cp_history:
                        with st.chat_message(msg["role"], avatar="🤖" if msg["role"] == "assistant" else None):
                            st.markdown(msg["content"])

                if cp_prompt:
                    db.save_cp_message(cp_id, "user", cp_prompt)

                    recent_cp = db.get_cp_chat(cp_id, limit=20)
                    _raw_cp = [
                        {"role": m["role"], "content": m["content"]}
                        for m in recent_cp[:-1]
                        if (m.get("content") or "").strip()
                    ]
                    cp_api_msgs = _dedup_chat(_raw_cp)
                    cp_api_msgs.append({"role": "user", "content": cp_prompt})

                    cp_final = ""
                    with _cp_chat_container:
                        with st.chat_message("user"):
                            st.markdown(cp_prompt)
                        with st.chat_message("assistant", avatar="🤖"):
                            with st.spinner("Responding…"):
                                try:
                                    _cp_client = anthropic.Anthropic(api_key=api_key)
                                    # No extended thinking for chat — keeps latency low
                                    _cp_resp = _cp_client.messages.create(
                                        model=_cp_chat_model,
                                        max_tokens=2048,
                                        system=_cp_chat_system,
                                        messages=cp_api_msgs,
                                    )
                                    for block in _cp_resp.content:
                                        if hasattr(block, "text"):
                                            cp_final += block.text
                                except Exception as exc:
                                    cp_final = f"⚠️ Error ({type(exc).__name__}): {exc}"
                            st.markdown(cp_final)

                    if not cp_final.strip():
                        cp_final = "(no text response received)"
                    db.save_cp_message(cp_id, "assistant", cp_final)
                    st.rerun()


def _opp_inline_panel(opp_id: int):
    """Inline panel with Edit / Scope / Proposal / Chat tabs for an opportunity."""
    opp = db.get_opportunity(opp_id)
    if not opp:
        st.error("Opportunity not found.")
        if st.button("✖ Close", key=f"panel_close_err_{opp_id}"):
            st.session_state.pop("open_opp_id", None)
            st.rerun()
        return

    _stage_c = STAGE_COLOR.get(opp["stage"], "#888")
    _stage_bg = STAGE_BG.get(opp["stage"], "#F5F5F5")

    with st.container(border=True):
        _ph1, _ph2 = st.columns([9, 1])
        _ph1.markdown(
            f"<h4 style='margin:0;color:#33475B'>"
            f"{opp.get('company','')} — {opp.get('title','')}</h4>"
            f"<span style='background:{_stage_bg};color:{_stage_c};padding:2px 10px;"
            f"border-radius:10px;font-size:12px;font-weight:700'>{opp['stage']}</span>"
            f"&nbsp;&nbsp;<span style='color:#7C98B6;font-size:13px'>"
            f"S${opp.get('value_sgd',0):,.0f} &nbsp;·&nbsp; "
            f"EP: {opp.get('engagement_partner','—')} &nbsp;·&nbsp; "
            f"EM: {opp.get('engagement_manager','—')}</span>",
            unsafe_allow_html=True,
        )
        if _ph2.button("✖ Close", key=f"panel_close_{opp_id}", use_container_width=True):
            st.session_state.pop("open_opp_id", None)
            st.rerun()

        # Fetch shared data once — reused across all tabs
        _all_clients = db.get_clients()
        _scope_docs = db.get_opportunity_files(opp_id)
        _existing_prop = db.get_proposal(opp_id)

        tab_edit, tab_scope, tab_proposal, tab_chat = st.tabs(
            ["✏️ Edit", "📁 Scope", "📄 Proposal", "💬 Chat with Aishah"]
        )

        # ══ TAB 1 — EDIT ══════════════════════════════════════════════════════
        with tab_edit:
            _deal_form(opp, _all_clients, form_key=f"opp_form_{opp_id}")

        # ══ TAB 2 — SCOPE FILES ═══════════════════════════════════════════════
        with tab_scope:
            st.markdown("##### Scope documents (RFP, TOR, briefing decks, emails)")

            # ── Paste text directly ───────────────────────────────────────────
            with st.expander("📋 Paste text directly", expanded=False):
                _paste_text = st.text_area(
                    "Paste content here",
                    height=180,
                    placeholder="Paste any relevant text here…",
                    key=f"paste_text_{opp_id}",
                    label_visibility="collapsed",
                )
                if st.button("💾 Save as Text Summary", key=f"paste_save_{opp_id}", type="primary"):
                    _ptxt = _paste_text.strip()
                    _pname = "Text Summary"
                    if _ptxt:
                        db.add_opportunity_file({
                            "opportunity_id": opp_id,
                            "filename": _pname,
                            "file_type": "TXT",
                            "extracted_text": _ptxt,
                            "file_size_kb": round(len(_ptxt) / 1024, 1),
                        })
                        st.success(f"Saved '{_pname}' — {len(_ptxt):,} characters.")
                        st.rerun()
                    else:
                        st.warning("Nothing to save — paste some text first.")

            st.caption("Or upload a file: PDF, DOCX, TXT, Excel (XLSX/XLS).")
            uploaded = st.file_uploader(
                "Drop files here", type=["pdf", "docx", "txt", "xlsx", "xls"],
                accept_multiple_files=True, label_visibility="collapsed",
                key=f"scope_up_{opp_id}",
            )
            if uploaded:
                existing_names = {doc["filename"] for doc in _scope_docs}
                new_files = [uf for uf in uploaded if uf.name not in existing_names]
                skipped = [uf.name for uf in uploaded if uf.name in existing_names]
                if skipped:
                    st.info(f"Already on file (skipped): {', '.join(skipped)}")
                if new_files:
                    for uf in new_files:
                        with st.spinner(f"Extracting text from {uf.name}…"):
                            text = _extract_text(uf)
                        db.add_opportunity_file({
                            "opportunity_id": opp_id,
                            "filename": uf.name,
                            "file_type": uf.name.rsplit(".", 1)[-1].upper(),
                            "extracted_text": text,
                            "file_size_kb": round(uf.size / 1024, 1),
                        })
                    st.success(f"Added {len(new_files)} new file(s).")
                    st.rerun()

            scope_docs = _scope_docs
            # Auto-remove duplicates on load (keeps most recent per filename)
            _filenames_seen: set = set()
            _dup_ids = []
            for _d in scope_docs:
                if _d["filename"] in _filenames_seen:
                    _dup_ids.append(_d["id"])
                else:
                    _filenames_seen.add(_d["filename"])
            if _dup_ids:
                for _did in _dup_ids:
                    db.delete_opportunity_file(_did)
                scope_docs = db.get_opportunity_files(opp_id)

            if scope_docs:
                _sf1, _sf2 = st.columns([6, 2])
                _sf1.markdown(f"**{len(scope_docs)} document(s) on file:**")
                for doc in scope_docs:
                    _dc1, _dc2, _dc3 = st.columns([6, 2, 1])
                    _dc1.markdown(f"📎 **{doc['filename']}**")
                    _dc2.caption(f"{doc.get('file_size_kb', 0):.1f} KB · {doc.get('file_type','')} · {(doc.get('created_at') or '')[:10]}")
                    if _dc3.button("🗑️", key=f"del_file_{doc['id']}", help="Delete"):
                        db.delete_opportunity_file(doc["id"])
                        st.rerun()
                    with st.expander(f"Preview — {doc['filename']}", expanded=False):
                        preview = (doc.get("extracted_text") or "")[:1500]
                        st.text(preview + ("…" if len(doc.get("extracted_text") or "") > 1500 else ""))
            else:
                st.info("No documents uploaded yet. Upload scope files above, then generate a proposal.")

        # ══ TAB 2 — PROPOSAL ══════════════════════════════════════════════════════
        with tab_proposal:
            scope_docs = _scope_docs
            existing_prop = _existing_prop
            proposal_data = None

            if existing_prop:
                try:
                    proposal_data = json.loads(existing_prop["content"])
                except Exception:
                    proposal_data = None

            _pg1, _pg2, _pg3 = st.columns([3, 2, 2])
            api_configured = bool(os.environ.get("ANTHROPIC_API_KEY", ""))

            if _pg1.button(
                "✨ Generate Proposal" if not existing_prop else "🔄 Regenerate",
                type="primary", use_container_width=True,
                disabled=not api_configured,
                help="Uses Claude to generate a proposal from your scope documents and deal info",
            ):
                with st.spinner("Generating proposal with Claude…"):
                    new_prop, _gen_err = _generate_proposal(opp, scope_docs)
                if new_prop:
                    db.upsert_proposal(opp_id, json.dumps(new_prop))
                    st.success("Proposal generated.")
                    st.rerun()
                else:
                    st.error(f"Generation failed: {_gen_err}")

            if existing_prop and _pg2.button("🗑️ Delete Proposal", use_container_width=True):
                db.delete_proposal(opp_id)
                st.rerun()

            if not api_configured:
                st.warning("Add `ANTHROPIC_API_KEY` to Streamlit secrets to enable proposal generation.")

            # ── Upload a prepared proposal for Aishah to use as context ───────────
            with st.expander("📤 Upload Prepared Proposal", expanded=not existing_prop):
                st.caption(
                    "Upload an existing proposal document (PDF, DOCX, TXT, or Excel). "
                    "Aishah will use it as context in the Chat tab."
                )
                _uploaded_prop = st.file_uploader(
                    "Proposal document",
                    type=["pdf", "docx", "txt", "xlsx", "xls"],
                    key=f"prop_upload_{opp_id}",
                    label_visibility="collapsed",
                )
                if _uploaded_prop:
                    if st.button("📥 Save as proposal context", key=f"prop_upload_save_{opp_id}", type="primary"):
                        with st.spinner(f"Extracting text from {_uploaded_prop.name}…"):
                            _prop_text = _extract_text(_uploaded_prop)
                        if _prop_text.strip():
                            db.upsert_proposal(opp_id, _prop_text)
                            st.success(f"Saved — {len(_prop_text):,} characters extracted from {_uploaded_prop.name}.")
                            st.rerun()
                        else:
                            st.error("Could not extract text from that file. Try a different format.")

            # If content exists but isn't JSON (i.e. an uploaded document), show it as plain text
            if existing_prop and not proposal_data:
                st.divider()
                st.caption("📄 Uploaded proposal document — Aishah uses this as context in the Chat tab.")
                st.text_area(
                    "Proposal content",
                    value=existing_prop["content"],
                    height=300,
                    key=f"prop_plain_{opp_id}",
                    disabled=True,
                    label_visibility="collapsed",
                )

            if proposal_data:
                st.divider()

                # ── Executive Summary & Scope ─────────────────────────────────────
                with st.expander("📝 Executive Summary & Scope Understanding", expanded=True):
                    _new_exec = st.text_area(
                        "Executive Summary", value=proposal_data.get("executive_summary", ""),
                        height=140, key=f"exec_sum_{opp_id}",
                    )
                    _new_scope = st.text_area(
                        "Scope Understanding", value=proposal_data.get("scope_understanding", ""),
                        height=140, key=f"scope_und_{opp_id}",
                    )
                    if st.button("💾 Save", key=f"save_exec_{opp_id}"):
                        proposal_data["executive_summary"] = _new_exec
                        proposal_data["scope_understanding"] = _new_scope
                        db.upsert_proposal(opp_id, json.dumps(proposal_data))
                        st.success("Saved.")

                # ── Activities ────────────────────────────────────────────────────
                with st.expander("🗂️ Project Activities & Phases", expanded=True):
                    _acts = proposal_data.get("activities", [])
                    _act_rows = []
                    for ph in _acts:
                        for act in ph.get("activities", []):
                            _act_rows.append({
                                "Phase": ph.get("phase", ""),
                                "Activity": act,
                                "Duration (wks)": ph.get("duration_weeks", 0),
                                "Deliverable": ", ".join(ph.get("deliverables", [])),
                            })
                    if _act_rows:
                        _act_df = pd.DataFrame(_act_rows)
                        _act_edited = st.data_editor(
                            _act_df, use_container_width=True, hide_index=True,
                            num_rows="dynamic",
                            column_config={
                                "Phase": st.column_config.TextColumn("Phase"),
                                "Activity": st.column_config.TextColumn("Activity", width="large"),
                                "Duration (wks)": st.column_config.NumberColumn("Duration (wks)", min_value=0),
                                "Deliverable": st.column_config.TextColumn("Deliverable"),
                            },
                            key=f"act_editor_{opp_id}",
                        )
                        if st.button("💾 Save Activities", key=f"save_acts_{opp_id}"):
                            # Rebuild activities from edited flat rows
                            phases_map: dict = {}
                            for _, row in _act_edited.iterrows():
                                ph_name = str(row["Phase"]) if row["Phase"] else "General"
                                if ph_name not in phases_map:
                                    phases_map[ph_name] = {
                                        "phase": ph_name, "description": "",
                                        "activities": [], "duration_weeks": int(row["Duration (wks)"] or 0),
                                        "deliverables": [],
                                    }
                                if row["Activity"]:
                                    phases_map[ph_name]["activities"].append(str(row["Activity"]))
                                if row["Deliverable"]:
                                    for d in str(row["Deliverable"]).split(","):
                                        d = d.strip()
                                        if d and d not in phases_map[ph_name]["deliverables"]:
                                            phases_map[ph_name]["deliverables"].append(d)
                            proposal_data["activities"] = list(phases_map.values())
                            db.upsert_proposal(opp_id, json.dumps(proposal_data))
                            st.success("Activities saved.")
                    else:
                        st.info("No activities in this proposal yet.")

                # ── Resources ─────────────────────────────────────────────────────
                with st.expander("👥 Resources & Rates", expanded=True):
                    _res = proposal_data.get("resources", [])
                    _res_rows = []
                    for r in _res:
                        _res_rows.append({
                            "Role": r.get("role", ""),
                            "Rate/Day (S$)": float(r.get("rate_per_day", 0)),
                            "Days": int(r.get("days", 0)),
                            "Total (S$)": float(r.get("rate_per_day", 0)) * int(r.get("days", 0)),
                        })
                    if not _res_rows:
                        _res_rows = [{"Role": "", "Rate/Day (S$)": 0.0, "Days": 0, "Total (S$)": 0.0}]
                    _res_df = pd.DataFrame(_res_rows)
                    _res_edited = st.data_editor(
                        _res_df, use_container_width=True, hide_index=True,
                        num_rows="dynamic",
                        column_config={
                            "Role": st.column_config.SelectboxColumn("Role", options=_PWC_ROLES),
                            "Rate/Day (S$)": st.column_config.NumberColumn("Rate/Day (S$)", format="S$%,.0f", min_value=0),
                            "Days": st.column_config.NumberColumn("Days", format="%d", min_value=0),
                            "Total (S$)": st.column_config.NumberColumn("Total (S$)", format="S$%,.0f", disabled=True),
                        },
                        key=f"res_editor_{opp_id}",
                    )
                    # Auto-fill rate when role is selected
                    _total_fees = sum(
                        float(r["Rate/Day (S$)"] or 0) * int(r["Days"] or 0)
                        for _, r in _res_edited.iterrows()
                    )
                    st.metric("Total Fees", f"S${_total_fees:,.0f}")
                    if st.button("💾 Save Resources", key=f"save_res_{opp_id}"):
                        proposal_data["resources"] = [
                            {"role": str(r["Role"]) if r["Role"] else "",
                             "rate_per_day": float(r["Rate/Day (S$)"] or 0),
                             "days": int(r["Days"] or 0)}
                            for _, r in _res_edited.iterrows() if r["Role"]
                        ]
                        # Recalculate pricing breakdown from resources
                        proposal_data["pricing_breakdown"] = [
                            {"item": str(r["Role"]),
                             "amount": float(r["Rate/Day (S$)"] or 0) * int(r["Days"] or 0),
                             "notes": f"{r['Days']} day(s) × S${r['Rate/Day (S$)']:,.0f}"}
                            for _, r in _res_edited.iterrows() if r["Role"]
                        ]
                        db.upsert_proposal(opp_id, json.dumps(proposal_data))
                        st.success("Resources saved.")

                # ── Timeline / Milestones ─────────────────────────────────────────
                with st.expander("📅 Timeline & Milestones", expanded=False):
                    _ms = proposal_data.get("milestones", [])
                    _ms_rows = [{"Week": int(m.get("week", 0)), "Milestone": m.get("milestone", "")}
                                for m in _ms]
                    if not _ms_rows:
                        _ms_rows = [{"Week": 1, "Milestone": ""}]
                    _ms_df = pd.DataFrame(_ms_rows)
                    _ms_edited = st.data_editor(
                        _ms_df, use_container_width=True, hide_index=True, num_rows="dynamic",
                        column_config={
                            "Week": st.column_config.NumberColumn("Week", min_value=1),
                            "Milestone": st.column_config.TextColumn("Milestone", width="large"),
                        },
                        key=f"ms_editor_{opp_id}",
                    )
                    if st.button("💾 Save Milestones", key=f"save_ms_{opp_id}"):
                        proposal_data["milestones"] = [
                            {"week": int(r["Week"] or 0), "milestone": str(r["Milestone"]) if r["Milestone"] else ""}
                            for _, r in _ms_edited.iterrows() if r["Milestone"]
                        ]
                        db.upsert_proposal(opp_id, json.dumps(proposal_data))
                        st.success("Milestones saved.")

                # ── Pricing ───────────────────────────────────────────────────────
                with st.expander("💰 Pricing Summary", expanded=True):
                    _pb = proposal_data.get("pricing_breakdown", [])
                    _pb_rows = [{"Item": p.get("item", ""), "Amount (S$)": float(p.get("amount", 0)),
                                 "Notes": p.get("notes", "")} for p in _pb]
                    if not _pb_rows:
                        _pb_rows = [{"Item": "", "Amount (S$)": 0.0, "Notes": ""}]
                    _pb_df = pd.DataFrame(_pb_rows)
                    _pb_edited = st.data_editor(
                        _pb_df, use_container_width=True, hide_index=True, num_rows="dynamic",
                        column_config={
                            "Item": st.column_config.TextColumn("Item"),
                            "Amount (S$)": st.column_config.NumberColumn("Amount (S$)", format="S$%,.0f", min_value=0),
                            "Notes": st.column_config.TextColumn("Notes"),
                        },
                        key=f"pb_editor_{opp_id}",
                    )
                    _grand_total = sum(float(r["Amount (S$)"] or 0) for _, r in _pb_edited.iterrows())
                    _deal_val = float(opp.get("value_sgd") or 0)
                    _diff = _grand_total - _deal_val
                    _mc1, _mc2, _mc3 = st.columns(3)
                    _mc1.metric("Proposal Total", f"S${_grand_total:,.0f}")
                    _mc2.metric("Deal Value", f"S${_deal_val:,.0f}")
                    _mc3.metric("Variance", f"S${_diff:+,.0f}",
                                delta_color="off" if abs(_diff) < 1 else "inverse" if _diff < 0 else "normal")
                    if st.button("💾 Save Pricing", key=f"save_pb_{opp_id}"):
                        proposal_data["pricing_breakdown"] = [
                            {"item": str(r["Item"]) if r["Item"] else "",
                             "amount": float(r["Amount (S$)"] or 0),
                             "notes": str(r["Notes"]) if r["Notes"] else ""}
                            for _, r in _pb_edited.iterrows() if r["Item"]
                        ]
                        db.upsert_proposal(opp_id, json.dumps(proposal_data))
                        st.success("Pricing saved.")

                # ── Assumptions & Exclusions ──────────────────────────────────────
                with st.expander("📋 Assumptions & Exclusions", expanded=False):
                    _ae1, _ae2 = st.columns(2)
                    _new_assumptions = _ae1.text_area(
                        "Assumptions (one per line)",
                        value="\n".join(proposal_data.get("assumptions", [])),
                        height=180, key=f"assum_{opp_id}",
                    )
                    _new_exclusions = _ae2.text_area(
                        "Exclusions (one per line)",
                        value="\n".join(proposal_data.get("exclusions", [])),
                        height=180, key=f"excl_{opp_id}",
                    )
                    if st.button("💾 Save", key=f"save_ae_{opp_id}"):
                        proposal_data["assumptions"] = [l.strip() for l in _new_assumptions.split("\n") if l.strip()]
                        proposal_data["exclusions"] = [l.strip() for l in _new_exclusions.split("\n") if l.strip()]
                        db.upsert_proposal(opp_id, json.dumps(proposal_data))
                        st.success("Saved.")

        # ══ TAB 3 — CHAT ══════════════════════════════════════════════════════════
        with tab_chat:
            api_key = os.environ.get("ANTHROPIC_API_KEY", "")
            if not api_key:
                st.warning("Add `ANTHROPIC_API_KEY` to Streamlit secrets to use the chat.")
            else:
                scope_docs = _scope_docs
                existing_prop = _existing_prop

                # Build context block shown to Claude on every turn
                _scope_ctx = "\n\n".join(
                    f"FILE: {d['filename']}\n{(d.get('extracted_text') or '')[:4000]}"
                    for d in scope_docs
                ) or "No scope documents uploaded."
                _prop_ctx = existing_prop["content"] if existing_prop else "No proposal generated yet."
                _opp_ctx = (
                    f"Opportunity: {opp.get('title','')} | Client: {opp.get('company','')}\n"
                    f"Value: S${opp.get('value_sgd',0):,.0f} | Stage: {opp['stage']}\n"
                    f"Service: {opp.get('ai_service_type','')} | EP: {opp.get('engagement_partner','')} EM: {opp.get('engagement_manager','')}\n"
                    f"Close Date: {opp.get('expected_close_date','TBD')}"
                )
                _chat_system = (
                    f"{_PROPOSAL_SYSTEM}\n\n"
                    f"=== CURRENT OPPORTUNITY ===\n{_opp_ctx}\n\n"
                    f"=== SCOPE DOCUMENTS ===\n{_scope_ctx[:6000]}\n\n"
                    f"=== CURRENT PROPOSAL ===\n{_prop_ctx[:6000]}\n\n"
                    "=== CHAT MODE — OVERRIDE ALL PREVIOUS OUTPUT INSTRUCTIONS ===\n"
                    "You are in CONVERSATIONAL mode. Rules (follow strictly):\n"
                    "1. For questions, analysis, pricing checks, status summaries, or any general discussion: "
                    "respond in plain English prose ONLY. No JSON of any kind.\n"
                    "2. ONLY return JSON when the user explicitly says words like 'generate', 'create', "
                    "'update', 'revise', 'draft', or 'write' the proposal document.\n"
                    "3. When you DO return a proposal JSON: write a 1-2 sentence summary first, then the "
                    "JSON wrapped in ```json\\n{...}\\n``` fences. The system silently saves the JSON — "
                    "the user only sees your prose summary.\n"
                    "4. NEVER return bare JSON without fences. NEVER return ONLY JSON with no prose.\n"
                    "5. When in doubt: respond in plain text."
                )

                history = _dedup_chat(db.get_opp_chat(opp_id, limit=40))
                _clr_col, _dbg_col, _opp_mdl_col, _ = st.columns([1, 1, 2, 3])
                if _clr_col.button("🗑️ Clear chat", key=f"clr_opp_chat_{opp_id}"):
                    db.clear_opp_chat(opp_id)
                    st.rerun()
                _show_debug = _dbg_col.toggle("🔍 Debug", key=f"opp_chat_debug_{opp_id}")
                _opp_chat_model = _CHAT_MODELS[_opp_mdl_col.selectbox(
                    "Model", list(_CHAT_MODELS.keys()),
                    index=list(_CHAT_MODELS.keys()).index(_CHAT_MODEL_DEFAULT),
                    key=f"opp_chat_model_{opp_id}", label_visibility="collapsed",
                )]

                if _show_debug:
                    with st.expander("🔍 Debug — chat state", expanded=True):
                        st.write(f"**opp_id:** {opp_id}")
                        st.write(f"**History rows in DB:** {len(history)}")
                        st.write(f"**System prompt length:** {len(_chat_system):,} chars")
                        st.write(f"**Scope ctx length:** {len(_scope_ctx):,} chars")
                        st.write(f"**Proposal ctx length:** {len(_prop_ctx):,} chars")
                        for i, m in enumerate(history):
                            role = m.get("role", "?")
                            content = m.get("content") or ""
                            st.write(f"  msg[{i}] role={role} len={len(content)} "
                                     f"empty={not content.strip()} preview={content[:80]!r}")

                _chat_container = st.container(height=420)
                with _chat_container:
                    for msg in history:
                        with st.chat_message(msg["role"], avatar="🤖" if msg["role"] == "assistant" else None):
                            st.markdown(msg["content"])

                prompt = st.chat_input(
                    "Ask Aishah — refine the proposal, adjust resources, add activities, price check…",
                    key=f"opp_chat_input_{opp_id}",
                )
                if prompt:
                    # Save user message and show spinner inside container immediately
                    db.save_opp_message(opp_id, "user", prompt)

                    recent = db.get_opp_chat(opp_id, limit=20)
                    # Build API messages: skip empties, deduplicate consecutive same-role
                    _raw_msgs = [
                        {"role": m["role"], "content": m["content"]}
                        for m in recent[:-1]
                        if (m.get("content") or "").strip()
                    ]
                    api_messages = _dedup_chat(_raw_msgs)
                    api_messages.append({"role": "user", "content": prompt})

                    if _show_debug:
                        st.info(f"DEBUG: sending {len(api_messages)} messages to API, "
                                f"total content chars: {sum(len(m['content']) for m in api_messages):,}, "
                                f"system chars: {len(_chat_system):,}")

                    final_text = ""
                    _stop_reason = ""
                    _api_error = ""
                    _block_types = []

                    # Show live spinner in the container while API call runs
                    with _chat_container:
                        with st.chat_message("user"):
                            st.markdown(prompt)
                        with st.chat_message("assistant", avatar="🤖"):
                            with st.spinner("Responding…"):
                                try:
                                    client_ai = anthropic.Anthropic(api_key=api_key)
                                    # No extended thinking for chat — keeps latency low
                                    resp = client_ai.messages.create(
                                        model=_opp_chat_model,
                                        max_tokens=4096,
                                        system=_chat_system,
                                        messages=api_messages,
                                    )
                                    _stop_reason = resp.stop_reason
                                    for block in resp.content:
                                        _block_types.append(type(block).__name__)
                                        if hasattr(block, "text"):
                                            final_text += block.text
                                except Exception as exc:
                                    _api_error = traceback.format_exc()
                                    final_text = f"⚠️ Error ({type(exc).__name__}): {exc}"

                            # Strip proposal JSON before rendering.
                            # Only treat a JSON block as a proposal update if it contains
                            # at least one known proposal key — prevents code examples from
                            # being silently swallowed.
                            _PROP_KEYS = ("executive_summary", "activities", "resources",
                                          "scope_understanding", "investment_summary")
                            _jm = re.search(r'```json\s*([\s\S]*?)\s*```', final_text)
                            _raw_json_str = None
                            if _jm:
                                _candidate = _jm.group(1).strip()
                                try:
                                    _candidate_parsed = json.loads(_candidate)
                                    if isinstance(_candidate_parsed, dict) and any(
                                        k in _candidate_parsed for k in _PROP_KEYS
                                    ):
                                        _raw_json_str = _candidate
                                        _jm_start, _jm_end = _jm.start(), _jm.end()
                                except Exception:
                                    pass  # Invalid or non-proposal JSON in fence — leave as-is
                            if _raw_json_str is None:
                                # Non-greedy search for {...} blocks so we don't swallow prose
                                for _rm in re.finditer(r'\{[^{}]*(?:\{[^{}]*\}[^{}]*)+"?\}|\{[^{}]+\}', final_text, re.DOTALL):
                                    try:
                                        _probe = json.loads(_rm.group(0))
                                        if isinstance(_probe, dict) and any(
                                            k in _probe for k in _PROP_KEYS
                                        ):
                                            _raw_json_str = _rm.group(0)
                                            _jm_start, _jm_end = _rm.start(), _rm.end()
                                            break
                                    except Exception:
                                        continue

                            if _raw_json_str is not None:
                                _before = final_text[:_jm_start].strip()
                                _after = final_text[_jm_end:].strip()
                                try:
                                    _updated_prop = json.loads(_raw_json_str)
                                    db.upsert_proposal(opp_id, json.dumps(_updated_prop))
                                    final_text = (_before + ("\n\n" if _before else "") + _after).strip()
                                    if not final_text:
                                        final_text = "✅ Proposal updated — switch to the Proposal tab to review changes."
                                    else:
                                        final_text += "\n\n✅ Proposal updated automatically."
                                except Exception:
                                    # JSON parse failed (e.g. truncated by token limit) —
                                    # keep original text so the user always sees something
                                    pass
                            st.markdown(final_text)

                    if _show_debug:
                        st.write(f"**stop_reason:** {_stop_reason}")
                        st.write(f"**block types:** {_block_types}")
                        st.write(f"**final_text length:** {len(final_text)}")
                        if _api_error:
                            st.code(_api_error, language="text")

                    # Guard: never save empty assistant messages (would break next turn)
                    if not final_text.strip():
                        final_text = "(no text response — check Debug for block types)"

                    db.save_opp_message(opp_id, "assistant", final_text)
                    # Rerun so history container refreshes with the new messages at bottom
                    st.rerun()


# ─────────────────────────────────────────────────────────────────────────────
# DASHBOARD
# ─────────────────────────────────────────────────────────────────────────────
if page == "Dashboard":
    st.markdown(
        "<h1 style='font-size:2.4rem;font-weight:800;color:#33475B;letter-spacing:-0.5px;margin-bottom:0'>AISHAH</h1>"
        "<p style='color:#7C98B6;margin-top:2px;margin-bottom:16px'>AI Smart Hustles and Alliances Hub</p>",
        unsafe_allow_html=True,
    )
    all_opps = db.get_opportunities()
    active_opps = [o for o in all_opps if o["stage"] not in ("Won", "Lost")]
    won_opps = [o for o in all_opps if o["stage"] == "Won"]
    lost_opps = [o for o in all_opps if o["stage"] == "Lost"]
    total_pipeline = sum(o["value_sgd"] or 0 for o in active_opps)
    weighted = sum((o["value_sgd"] or 0) * (o["probability"] or 0) / 100 for o in active_opps)
    won_val = sum(o["value_sgd"] or 0 for o in won_opps)
    _closed = len(won_opps) + len(lost_opps)
    win_rate = round(len(won_opps) / _closed * 100) if _closed else 0
    avg_deal = total_pipeline / len(active_opps) if active_opps else 0
    all_activities = db.get_engagements(limit=200)
    overdue = [
        o for o in active_opps
        if o.get("next_action_date") and (_days_until(o["next_action_date"]) or 0) < 0
    ]

    if overdue:
        st.markdown(
            f"<div class='hs-overdue'>⚠️ <b>{len(overdue)} deal(s) have overdue follow-up actions</b> — "
            + ", ".join(o["company"] for o in overdue[:3])
            + ("..." if len(overdue) > 3 else "") + "</div>",
            unsafe_allow_html=True,
        )

    if "edit_opp" in st.session_state:
        with st.container(border=True):
            _deal_form(st.session_state["edit_opp"], db.get_clients(), form_key="opp_form_dash_new")

    # ── Active Opportunities ───────────────────────────────────────────────────
    # (inline panel is rendered by _deals_table when a row is tapped)
    _dl1, _dl2 = st.columns([5, 1])
    _dl1.subheader("🗂️ Active Opportunities")
    if _dl2.button("＋ New Opportunity", type="primary", key="dash_new_deal2", use_container_width=True):
        st.session_state["edit_opp"] = {}
        st.rerun()

    if active_opps:
        _deals_table(active_opps, key_prefix="dash_act", mode="active")
    else:
        st.info("No active opportunities yet — click '＋ New Opportunity' to create your first one.")

    # ── Won Opportunities ──────────────────────────────────────────────────────
    st.divider()
    _won_color_w = STAGE_COLOR["Won"]
    st.markdown(
        f"<h3 style='margin-bottom:4px'>✅ Won Opportunities "
        f"<span style='font-size:16px;color:#7C98B6;font-weight:400'>"
        f"{len(won_opps)} {'opportunity' if len(won_opps) == 1 else 'opportunities'} · {_fmt_sgd(won_val)}"
        f"</span></h3>",
        unsafe_allow_html=True,
    )

    if won_opps:
        _deals_table(won_opps, key_prefix="dash_won", mode="won")
    else:
        st.info("No won deals yet — keep pushing!")

    # ── Analytics ──────────────────────────────────────────────────────────────
    st.divider()
    col_l, col_r = st.columns([3, 2])

    with col_l:
        st.subheader("Pipeline by Stage")
        pipeline_summary = db.get_pipeline_summary()
        ps_df = {row["stage"]: row for row in pipeline_summary}
        funnel_stages = [s for s in STAGE_ORDER if s != "Lost"]
        funnel_vals = [ps_df.get(s, {}).get("total_value", 0) or 0 for s in funnel_stages]
        funnel_counts = [ps_df.get(s, {}).get("count", 0) or 0 for s in funnel_stages]
        funnel_labels = [
            f"{s}<br><small>{funnel_counts[i]} deals · {_fmt_sgd(funnel_vals[i])}</small>"
            for i, s in enumerate(funnel_stages)
        ]
        fig_funnel = go.Figure(go.Funnel(
            y=funnel_labels, x=funnel_vals, textinfo="none",
            marker_color=[STAGE_COLOR[s] for s in funnel_stages],
        ))
        fig_funnel.update_layout(
            height=300, margin=dict(l=0, r=0, t=10, b=0),
            paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
            font_color="#33475B",
        )
        st.plotly_chart(fig_funnel, use_container_width=True)

        st.subheader("Revenue Forecast by Close Month")
        monthly: dict = {}
        for o in active_opps:
            cd = o.get("expected_close_date", "")
            if cd:
                try:
                    mo = datetime.strptime(cd, "%Y-%m-%d").strftime("%Y-%m")
                    monthly.setdefault(mo, {"total": 0, "weighted": 0})
                    monthly[mo]["total"] += o["value_sgd"] or 0
                    monthly[mo]["weighted"] += (o["value_sgd"] or 0) * (o["probability"] or 0) / 100
                except ValueError:
                    pass
        if monthly:
            _months = sorted(monthly.keys())
            fig_fc = go.Figure()
            fig_fc.add_bar(x=_months, y=[monthly[m]["total"] for m in _months],
                           name="Full Value", marker_color="#5BA4E6", opacity=0.45)
            fig_fc.add_bar(x=_months, y=[monthly[m]["weighted"] for m in _months],
                           name="Weighted", marker_color="#FF7A59")
            fig_fc.update_layout(
                barmode="overlay", height=200, margin=dict(l=0, r=0, t=10, b=0),
                paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
                font_color="#33475B", legend=dict(orientation="h", y=1.1),
            )
            st.plotly_chart(fig_fc, use_container_width=True)
        else:
            st.caption("Set expected close dates on deals to see forecast.")

    with col_r:
        st.subheader("Actions Due — Next 14 Days")
        upcoming = db.get_upcoming_actions(14)
        if upcoming:
            for u in upcoming[:8]:
                d = _days_until(u["next_action_date"])
                badge = "🔴" if d is not None and d <= 3 else "🟡" if d is not None and d <= 7 else "🟢"
                buyer = "🏛️" if u["buyer_type"] == "Institutional" else "👤"
                st.markdown(
                    f"{badge} {buyer} **{u['company']}**  \n"
                    f"<small style='color:#7C98B6'>{u['next_action']} · {u['next_action_date']} · {u['stage']}</small>",
                    unsafe_allow_html=True,
                )
        else:
            st.info("No actions due in the next 14 days.")

        st.divider()
        st.subheader("Recent Activity")
        recent_acts = all_activities[:6]
        if recent_acts:
            for a in recent_acts:
                icon = ACTIVITY_ICONS.get(a["activity_type"], "📌")
                summary_snip = (a.get("summary") or "")[:60]
                st.markdown(
                    f"{icon} **{a['company']}** · {a['activity_type']}  \n"
                    f"<small style='color:#7C98B6'>{a['activity_date']}"
                    + (f" · {summary_snip}{'...' if len(a.get('summary',''))>60 else ''}" if summary_snip else "")
                    + "</small>",
                    unsafe_allow_html=True,
                )
        else:
            st.caption("No activities logged yet.")

        st.divider()
        wl1, wl2 = st.columns(2)
        wl1.metric("Won", f"{len(won_opps)} deals", delta=_fmt_sgd(won_val) if won_val else None)
        wl2.metric("Lost", str(len(lost_opps)))


# ─────────────────────────────────────────────────────────────────────────────
# PIPELINE
# ─────────────────────────────────────────────────────────────────────────────
elif page == "Pipeline":
    # Page header with quick-create
    _ph1, _ph2 = st.columns([6, 1])
    _ph1.markdown("# 💼 Pipeline")
    if _ph2.button("＋ New Opportunity", type="primary", use_container_width=True):
        st.session_state["edit_opp"] = {}

    all_opps = db.get_opportunities()
    active_opps = [o for o in all_opps if o["stage"] not in ("Won", "Lost")]
    won_opps = [o for o in all_opps if o["stage"] == "Won"]
    lost_opps = [o for o in all_opps if o["stage"] == "Lost"]
    total_pipeline = sum(o["value_sgd"] or 0 for o in active_opps)
    weighted_pipeline = sum((o["value_sgd"] or 0) * (o["probability"] or 0) / 100 for o in active_opps)
    won_value = sum(o["value_sgd"] or 0 for o in won_opps)
    _closed_count = len(won_opps) + len(lost_opps)
    win_rate = round(len(won_opps) / _closed_count * 100) if _closed_count else 0
    avg_deal = total_pipeline / len(active_opps) if active_opps else 0
    overdue = [
        o for o in active_opps
        if o.get("next_action_date") and (_days_until(o["next_action_date"]) or 0) < 0
    ]

    m1, m2, m3, m4, m5, m6 = st.columns(6)
    m1.metric("Total Pipeline", _fmt_sgd(total_pipeline))
    m2.metric("Weighted Forecast", _fmt_sgd(weighted_pipeline))
    m3.metric("Won (FY)", _fmt_sgd(won_value))
    m4.metric("Win Rate", f"{win_rate}%")
    m5.metric("Active Deals", len(active_opps))
    m6.metric("Avg Deal Size", _fmt_sgd(avg_deal))

    if overdue:
        st.markdown(
            f"<div class='hs-overdue'>⚠️ {len(overdue)} deal(s) have overdue follow-up actions — "
            + ", ".join(o["company"] for o in overdue[:3]) + "</div>",
            unsafe_allow_html=True,
        )

    if "edit_opp" in st.session_state:
        with st.container(border=True):
            _deal_form(st.session_state["edit_opp"], db.get_clients(), form_key="opp_form_pipeline_new")

    # Only show the top-of-page panel for Kanban-triggered opens.
    # When a table row is selected the panel renders inline inside the Table tab instead.
    _tbl_widget_state = st.session_state.get("pipeline_tbl_sel") or {}
    _tbl_has_row = bool((_tbl_widget_state.get("selection") or {}).get("rows"))
    if st.session_state.get("open_opp_id") and not _tbl_has_row:
        _opp_inline_panel(st.session_state["open_opp_id"])

    st.divider()

    tab_kanban, tab_table, tab_forecast = st.tabs(["🗂️ Kanban", "📋 Table", "📈 Forecast"])

    # ── Kanban ────────────────────────────────────────────────────────────────
    with tab_kanban:
        active_stages = [s for s in STAGE_ORDER if s not in ("Won", "Lost")]
        cols = st.columns(len(active_stages))

        for col_idx, stage in enumerate(active_stages):
            stage_deals = [o for o in active_opps if o["stage"] == stage]
            stage_total = sum(o["value_sgd"] or 0 for o in stage_deals)
            color = STAGE_COLOR[stage]
            bg = STAGE_BG[stage]

            next_stage_idx = STAGE_ORDER.index(stage) + 1
            next_stage = STAGE_ORDER[next_stage_idx] if next_stage_idx < len(STAGE_ORDER) else None

            with cols[col_idx]:
                # Column header
                st.markdown(
                    f"<div style='background:{bg};border-left:4px solid {color};"
                    f"border-radius:6px;padding:10px 12px;margin-bottom:12px'>"
                    f"<span style='font-size:12px;font-weight:700;color:{color};text-transform:uppercase;"
                    f"letter-spacing:0.5px'>{stage}</span><br>"
                    f"<span style='font-size:11px;color:#7C98B6'>{len(stage_deals)} deal{'s' if len(stage_deals)!=1 else ''}"
                    f" · {_fmt_sgd(stage_total)}</span></div>",
                    unsafe_allow_html=True,
                )

                if not stage_deals:
                    st.markdown(
                        "<div style='border:2px dashed #DFE3EB;border-radius:6px;padding:24px;"
                        "text-align:center;color:#B0C4D8;font-size:12px'>No deals</div>",
                        unsafe_allow_html=True,
                    )

                for o in stage_deals:
                    buyer_icon = "🏛️" if o.get("buyer_type") == "Institutional" else "👤"
                    close_b = _close_badge(o.get("expected_close_date", ""))
                    action_b = _action_badge(o.get("next_action_date", ""))
                    prob_color = "#00A862" if o["probability"] >= 70 else "#F5A623" if o["probability"] >= 40 else "#E74C3C"

                    st.markdown(
                        f"<div class='hs-deal' style='border-left:3px solid {color}'>"
                        f"<div style='font-size:10px;color:#7C98B6;font-weight:500;margin-bottom:2px'>"
                        f"{buyer_icon} {o['company']}</div>"
                        f"<div style='font-size:13px;font-weight:600;color:#33475B;line-height:1.3;"
                        f"margin-bottom:6px'>{o['title']}</div>"
                        f"<div style='display:flex;justify-content:space-between;align-items:center'>"
                        f"<span style='font-size:14px;font-weight:700;color:#33475B'>{_fmt_sgd(o['value_sgd'])}</span>"
                        f"<span style='font-size:11px;color:{prob_color};font-weight:700'>{o['probability']}%</span>"
                        f"</div>"
                        + (f"<div style='font-size:10px;color:#7C98B6;margin-top:4px'>📅 {close_b}</div>" if close_b else "")
                        + (f"<div style='font-size:10px;color:#7C98B6'>⚡ {o['next_action']} {action_b}</div>" if o.get("next_action") else "")
                        + "</div>",
                        unsafe_allow_html=True,
                    )

                    # Action buttons
                    if next_stage:
                        _b1, _b2, _b3 = st.columns(3)
                        _btn_label = f"→ {next_stage[:6]}"
                        if _b1.button(_btn_label, key=f"mv_{o['id']}", use_container_width=True,
                                      help=f"Move to {next_stage}"):
                            _opp = db.get_opportunity(o["id"])
                            _opp["stage"] = next_stage
                            db.upsert_opportunity(_opp)
                            st.rerun()
                        if _b2.button("✏️ Edit", key=f"ed_{o['id']}", use_container_width=True):
                            st.session_state["open_opp_id"] = o["id"]
                            st.session_state.pop("edit_opp", None)
                            st.rerun()
                        if _b3.button("🗑️", key=f"dl_{o['id']}", use_container_width=True,
                                      help="Delete deal"):
                            db.delete_opportunity(o["id"])
                            st.rerun()
                    else:
                        _b1, _b2 = st.columns(2)
                        if _b1.button("✏️ Edit", key=f"ed_{o['id']}", use_container_width=True):
                            st.session_state["open_opp_id"] = o["id"]
                            st.session_state.pop("edit_opp", None)
                            st.rerun()
                        if _b2.button("🗑️ Del", key=f"dl_{o['id']}", use_container_width=True):
                            db.delete_opportunity(o["id"])
                            st.rerun()

        # Won / Lost section
        st.divider()
        _ws1, _ws2 = st.columns(2)
        _won_color = STAGE_COLOR["Won"]
        _lost_color = STAGE_COLOR["Lost"]
        _lost_total = sum(o["value_sgd"] or 0 for o in lost_opps)

        with _ws1:
            st.markdown(
                f"<div style='background:{STAGE_BG['Won']};border-left:4px solid {_won_color};"
                f"border-radius:6px;padding:10px 14px;margin-bottom:10px'>"
                f"<span style='font-size:13px;font-weight:700;color:{_won_color}'>✅ Won</span>"
                f"<span style='font-size:12px;color:#7C98B6;margin-left:8px'>"
                f"{len(won_opps)} deals · {_fmt_sgd(won_value)}</span></div>",
                unsafe_allow_html=True,
            )
            for o in won_opps:
                st.markdown(
                    f"<div class='hs-won'><b>{o['company']}</b> — {o['title']}"
                    f"<span style='float:right;color:{_won_color};font-weight:700'>{_fmt_sgd(o['value_sgd'])}</span></div>",
                    unsafe_allow_html=True,
                )
                _we1, _we2 = st.columns([1, 1])
                if _we1.button("✏️ Edit", key=f"wo_ed_{o['id']}", use_container_width=True):
                    st.session_state["open_opp_id"] = o["id"]
                    st.session_state.pop("edit_opp", None)
                    st.rerun()
                if _we2.button("🗑️ Del", key=f"wo_dl_{o['id']}", use_container_width=True):
                    db.delete_opportunity(o["id"])
                    st.rerun()

        with _ws2:
            st.markdown(
                f"<div style='background:{STAGE_BG['Lost']};border-left:4px solid {_lost_color};"
                f"border-radius:6px;padding:10px 14px;margin-bottom:10px'>"
                f"<span style='font-size:13px;font-weight:700;color:{_lost_color}'>❌ Lost</span>"
                f"<span style='font-size:12px;color:#7C98B6;margin-left:8px'>"
                f"{len(lost_opps)} deals · {_fmt_sgd(_lost_total)}</span></div>",
                unsafe_allow_html=True,
            )
            for o in lost_opps:
                st.markdown(
                    f"<div class='hs-lost'><b>{o['company']}</b> — {o['title']}"
                    f"<span style='float:right'>{_fmt_sgd(o['value_sgd'])}</span></div>",
                    unsafe_allow_html=True,
                )
                _le1, _le2 = st.columns([1, 1])
                if _le1.button("✏️ Edit", key=f"lo_ed_{o['id']}", use_container_width=True):
                    st.session_state["open_opp_id"] = o["id"]
                    st.session_state.pop("edit_opp", None)
                    st.rerun()
                if _le2.button("🗑️ Del", key=f"lo_dl_{o['id']}", use_container_width=True):
                    db.delete_opportunity(o["id"])
                    st.rerun()

    # ── Table ─────────────────────────────────────────────────────────────────
    with tab_table:
        _tbl_stage = st.selectbox("Filter Stage", ["All"] + STAGE_ORDER, key="tbl_stage_sel")
        filtered = all_opps if _tbl_stage == "All" else [o for o in all_opps if o["stage"] == _tbl_stage]
        if filtered:
            _tdf = pd.DataFrame(filtered)
            _cols = ["company", "title", "stage", "value_sgd", "probability",
                     "ai_service_type", "contract_type", "expected_close_date",
                     "next_action_date", "next_action"]
            _avail = [c for c in _cols if c in _tdf.columns]
            _tdf_disp = _tdf[_avail].copy()
            _tdf_disp.columns = [c.replace("_", " ").title() for c in _avail]
            st.caption("Click any row to open that opportunity.")
            _tbl_sel = st.dataframe(
                _tdf_disp,
                use_container_width=True,
                hide_index=True,
                on_select="rerun",
                selection_mode="single-row",
                key="pipeline_tbl_sel",
            )
            _tbl_rows = _tbl_sel.selection.rows
            if _tbl_rows:
                _clicked_opp = filtered[_tbl_rows[0]]
                st.session_state["open_opp_id"] = _clicked_opp["id"]
                st.session_state.pop("edit_opp", None)
                st.divider()
                _opp_inline_panel(_clicked_opp["id"])
        else:
            st.info("No deals found.")

    # ── Forecast ──────────────────────────────────────────────────────────────
    with tab_forecast:
        st.subheader("Revenue Forecast by Close Month")
        monthly: dict = {}
        for o in active_opps:
            cd = o.get("expected_close_date", "")
            if cd:
                try:
                    mo = datetime.strptime(cd, "%Y-%m-%d").strftime("%Y-%m")
                    monthly.setdefault(mo, {"total": 0, "weighted": 0, "count": 0})
                    monthly[mo]["total"] += o["value_sgd"] or 0
                    monthly[mo]["weighted"] += (o["value_sgd"] or 0) * (o["probability"] or 0) / 100
                    monthly[mo]["count"] += 1
                except ValueError:
                    pass
        if monthly:
            _months = sorted(monthly.keys())
            fig_fc = go.Figure()
            fig_fc.add_bar(x=_months, y=[monthly[m]["total"] for m in _months],
                           name="Full Value", marker_color="#5BA4E6", opacity=0.45)
            fig_fc.add_bar(x=_months, y=[monthly[m]["weighted"] for m in _months],
                           name="Weighted", marker_color="#FF7A59")
            fig_fc.update_layout(
                barmode="overlay", height=300, margin=dict(l=0, r=0, t=20, b=0),
                paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
                font_color="#33475B", legend=dict(orientation="h"),
            )
            st.plotly_chart(fig_fc, use_container_width=True)
            _fc_df = pd.DataFrame([
                {"Month": m, "Deals": monthly[m]["count"],
                 "Total (S$)": _fmt_sgd(monthly[m]["total"]),
                 "Weighted (S$)": _fmt_sgd(monthly[m]["weighted"])}
                for m in _months
            ])
            st.dataframe(_fc_df, use_container_width=True, hide_index=True)
        else:
            st.info("Add expected close dates to deals to see the revenue forecast.")

        st.subheader("Upcoming Closes — Next 90 Days")
        _soon = sorted(
            [o for o in active_opps if o.get("expected_close_date") and
             0 <= (_days_until(o["expected_close_date"]) or 999) <= 90],
            key=lambda x: x.get("expected_close_date", ""),
        )
        if _soon:
            for o in _soon:
                d = _days_until(o["expected_close_date"])
                badge = f"🔴 {d}d" if d <= 14 else f"🟡 {d}d" if d <= 30 else f"🟢 {d}d"
                st.markdown(
                    f"{badge} · **{o['company']}** — {o['title']} · "
                    f"**{_fmt_sgd(o['value_sgd'])}** ({o['probability']}%) · {o['stage']}"
                )
        else:
            st.info("No deals closing in the next 90 days.")

        st.subheader("Overdue Follow-up Actions")
        if overdue:
            for o in sorted(overdue, key=lambda x: x.get("next_action_date", "")):
                d = _days_until(o["next_action_date"])
                st.markdown(
                    f"🔴 **{o['company']}** — {o['next_action']} "
                    f"(due {o['next_action_date']}, {abs(d)}d overdue) · {o['stage']}"
                )
        else:
            st.success("No overdue actions — you're on top of it!")


# ─────────────────────────────────────────────────────────────────────────────
# ACTIVE PROJECTS
# ─────────────────────────────────────────────────────────────────────────────
elif page == "Active Projects":
    st.markdown("# 🚧 Active Projects")

    _proj_all = db.get_opportunities()
    _projects = [o for o in _proj_all if o["stage"] == "Won"]

    if not _projects:
        st.info("No Won deals yet — projects appear here automatically when a deal is marked Won.")
    else:
        # Summary metrics
        _pr_contract = sum(o["value_sgd"] or 0 for o in _projects)
        _pr_pwc = sum(o.get("pwc_revenue") or 0 for o in _projects)
        _pr_wip = sum(o.get("wip") or 0 for o in _projects)
        _pr_non_pwc = sum(o.get("non_pwc_revenue") or 0 for o in _projects)
        _pr_rtg = sum(
            ((o["value_sgd"] or 0) - (o.get("non_pwc_revenue") or 0) - (o.get("wip") or 0))
            for o in _projects
        )
        _pm1, _pm2, _pm3, _pm4, _pm5 = st.columns(5)
        _pm1.metric("Contract Value", _fmt_sgd(_pr_contract))
        _pm2.metric("PwC Revenue", _fmt_sgd(_pr_pwc))
        _pm3.metric("Non-PwC Revenue", _fmt_sgd(_pr_non_pwc))
        _pm4.metric("WIP", _fmt_sgd(_pr_wip))
        _pm5.metric("Revenue to Go", _fmt_sgd(_pr_rtg))

        st.divider()

        # Inline edit form
        if "edit_proj" in st.session_state:
            _ep_o = st.session_state["edit_proj"]
            _ep_title = _ep_o.get("title", "")
            with st.expander(f"✏️ Edit Project — {_ep_title}", expanded=True):
                with st.form("proj_edit_form"):
                    _pf1, _pf2 = st.columns(2)
                    _pf_ep = _pf1.text_input("EP (Engagement Partner)", value=_ep_o.get("engagement_partner", ""))
                    _pf_em = _pf2.text_input("EM (Engagement Manager)", value=_ep_o.get("engagement_manager", ""))

                    _pf3, _pf4 = st.columns(2)
                    _pf_psd_val = None
                    if _ep_o.get("project_start_date"):
                        try:
                            _pf_psd_val = datetime.strptime(_ep_o["project_start_date"], "%Y-%m-%d").date()
                        except ValueError:
                            pass
                    _pf_ped_val = None
                    if _ep_o.get("project_end_date"):
                        try:
                            _pf_ped_val = datetime.strptime(_ep_o["project_end_date"], "%Y-%m-%d").date()
                        except ValueError:
                            pass
                    _pf_start = _pf3.date_input("Start Date", value=_pf_psd_val)
                    _pf_end = _pf4.date_input("End Date", value=_pf_ped_val)

                    _pf5, _pf6, _pf7, _pf8 = st.columns(4)
                    _pf_contract = _pf5.number_input(
                        "Contract Value (S$K)", min_value=0.0,
                        value=float(_ep_o.get("value_sgd") or 0) / 1000, step=50.0, format="%.0f",
                    )
                    _pf_pwc = _pf6.number_input(
                        "PwC Revenue", min_value=0.0,
                        value=float(_ep_o.get("pwc_revenue") or 0), step=10_000.0, format="%.0f",
                    )
                    _pf_non_pwc = _pf7.number_input(
                        "Non-PwC Revenue", min_value=0.0,
                        value=float(_ep_o.get("non_pwc_revenue") or 0), step=10_000.0, format="%.0f",
                    )
                    _pf_wip = _pf8.number_input(
                        "WIP", min_value=0.0,
                        value=float(_ep_o.get("wip") or 0), step=10_000.0, format="%.0f",
                    )
                    _rtg_preview = _pf_contract * 1000 - _pf_non_pwc - _pf_wip
                    st.caption(f"Revenue to Go (calculated): **{_fmt_sgd(_rtg_preview)}**")

                    _pf_notes = st.text_area("Notes", value=_ep_o.get("notes", ""), height=80)

                    _psv1, _psv2 = st.columns(2)
                    _psaved = _psv1.form_submit_button("💾 Save", type="primary", use_container_width=True)
                    _pcancelled = _psv2.form_submit_button("Cancel", use_container_width=True)

                    if _psaved:
                        db.upsert_opportunity({
                            **_ep_o,
                            "engagement_partner": _pf_ep,
                            "engagement_manager": _pf_em,
                            "project_start_date": str(_pf_start) if _pf_start else None,
                            "project_end_date": str(_pf_end) if _pf_end else None,
                            "value_sgd": _pf_contract * 1000,
                            "pwc_revenue": _pf_pwc,
                            "non_pwc_revenue": _pf_non_pwc,
                            "wip": _pf_wip,
                            "notes": _pf_notes,
                        })
                        del st.session_state["edit_proj"]
                        st.rerun()
                    if _pcancelled:
                        del st.session_state["edit_proj"]
                        st.rerun()

        # Project table sorted EP → EM → End Date
        _deals_table(_projects, key_prefix="proj", mode="projects")


# ─────────────────────────────────────────────────────────────────────────────
# ACCOUNT PLAN
# ─────────────────────────────────────────────────────────────────────────────
elif page == "Account Plan":
    st.markdown("# 📋 Account Plan")

    _ap_clients = db.get_clients()
    _ap_names = [c["company"] for c in _ap_clients]

    if not _ap_clients:
        st.info("No accounts yet — add a company first.")
    else:
        # Preserve selection across reruns
        _ap_default_idx = 0
        if st.session_state.get("ap_selected_company") in _ap_names:
            _ap_default_idx = _ap_names.index(st.session_state["ap_selected_company"])

        _ap_selected = st.selectbox(
            "Select Account", _ap_names, index=_ap_default_idx, label_visibility="collapsed",
            placeholder="Choose an account…",
        )
        st.session_state["ap_selected_company"] = _ap_selected
        _ap_c = next((c for c in _ap_clients if c["company"] == _ap_selected), None)

        if _ap_c:
            _ap_opps = db.get_opportunities(client_id=_ap_c["id"])
            _ap_active = [o for o in _ap_opps if o["stage"] not in ("Won", "Lost")]
            _ap_won = [o for o in _ap_opps if o["stage"] == "Won"]
            _ap_lost = [o for o in _ap_opps if o["stage"] == "Lost"]
            _ap_pipeline = sum(o["value_sgd"] or 0 for o in _ap_active)
            _ap_won_val = sum(o["value_sgd"] or 0 for o in _ap_won)
            _ap_acts = db.get_engagements(client_id=_ap_c["id"], limit=50)

            TIER_COLOR = {
                "Strategic": "#E8692A", "Key": "#0091AE",
                "Growth": "#00A862", "Monitor": "#9B9B9B",
            }
            _tier = _ap_c.get("account_tier") or ""
            _tier_html = (
                f"<span style='background:{TIER_COLOR.get(_tier,'#516F90')};color:white;"
                f"font-size:11px;font-weight:700;padding:2px 8px;border-radius:10px;"
                f"margin-left:10px'>{_tier}</span>"
                if _tier else ""
            )
            _rel_stars = "⭐" * (_ap_c.get("relationship_score") or 0)
            st.markdown(
                f"<h2 style='margin-bottom:2px'>{_ap_c['company']}{_tier_html}</h2>"
                f"<p style='color:#7C98B6;margin:0'>{_ap_c['sector']}"
                + (f" &nbsp;·&nbsp; {_ap_c.get('country','')}" if _ap_c.get("country") else "")
                + (f" &nbsp;·&nbsp; {_rel_stars}" if _rel_stars else "")
                + "</p>",
                unsafe_allow_html=True,
            )

            _ap_m1, _ap_m2, _ap_m3 = st.columns(3)
            _ap_m1.metric("Active Pipeline", _fmt_sgd(_ap_pipeline))
            _ap_m2.metric("Won Revenue", _fmt_sgd(_ap_won_val))
            _ap_m3.metric("Active Deals", len(_ap_active))

            st.divider()

            _tab_ov, _tab_si, _tab_sh, _tab_deals, _tab_ap, _tab_edit = st.tabs([
                "📄 Overview", "🔍 Strategic Intel", "👥 Stakeholders",
                "💼 Deals & Projects", "🎯 Account Plan", "✏️ Edit Plan",
            ])

            # ── Overview ──────────────────────────────────────────────────────
            with _tab_ov:
                _ov1, _ov2 = st.columns(2)
                with _ov1:
                    st.markdown("#### Company Profile")
                    _ov_rows = [
                        ("Sector", _ap_c.get("sector", "")),
                        ("Sub-sector", _ap_c.get("sub_sector", "")),
                        ("Annual Revenue", _ap_c.get("annual_revenue", "")),
                        ("Employees", _ap_c.get("employee_count", "")),
                        ("Company Size", _ap_c.get("company_size", "")),
                        ("Country", _ap_c.get("country", "")),
                        ("Website", _ap_c.get("website", "")),
                        ("AI Maturity", _ap_c.get("ai_maturity", "")),
                        ("Budget Cycle", _ap_c.get("budget_cycle", "")),
                    ]
                    for label, val in _ov_rows:
                        if val:
                            st.markdown(
                                f"<div style='display:flex;gap:8px;margin-bottom:4px'>"
                                f"<span style='color:#7C98B6;min-width:130px;font-size:13px'>{label}</span>"
                                f"<span style='color:#33475B;font-size:13px;font-weight:500'>{val}</span>"
                                f"</div>",
                                unsafe_allow_html=True,
                            )
                with _ov2:
                    st.markdown("#### Our Relationship")
                    _rel_rows = [
                        ("Buyer Type", _ap_c.get("buyer_type", "")),
                        ("Relationship Score", "⭐" * (_ap_c.get("relationship_score") or 0)),
                        ("Key Contact", _ap_c.get("key_contact", "")),
                        ("Contact Title", _ap_c.get("contact_title", "")),
                        ("Account Tier", _ap_c.get("account_tier", "")),
                        ("Our Exec Sponsor", _ap_c.get("executive_sponsor", "")),
                        ("Client Champion", _ap_c.get("champion", "")),
                    ]
                    for label, val in _rel_rows:
                        if val:
                            st.markdown(
                                f"<div style='display:flex;gap:8px;margin-bottom:4px'>"
                                f"<span style='color:#7C98B6;min-width:130px;font-size:13px'>{label}</span>"
                                f"<span style='color:#33475B;font-size:13px;font-weight:500'>{val}</span>"
                                f"</div>",
                                unsafe_allow_html=True,
                            )
                if _ap_c.get("notes"):
                    st.markdown("#### Notes")
                    st.markdown(
                        f"<div style='background:#F5F8FA;border-left:3px solid #CBD6E2;"
                        f"padding:10px 14px;border-radius:4px;font-size:13px;color:#33475B'>"
                        f"{_ap_c['notes']}</div>",
                        unsafe_allow_html=True,
                    )
                if _ap_acts:
                    st.markdown("#### Recent Engagement")
                    for a in _ap_acts[:4]:
                        icon = ACTIVITY_ICONS.get(a["activity_type"], "📌")
                        st.markdown(
                            f"{icon} **{a['activity_type']}** · {a['activity_date']}"
                            + (f" — {(a.get('summary') or '')[:100]}" if a.get("summary") else "")
                        )

            # ── Strategic Intel ───────────────────────────────────────────────
            with _tab_si:
                _si_sections = [
                    ("🏢 Company Strategy & Priorities", "business_strategy",
                     "What is the company focused on? Key strategic initiatives, growth areas, transformation agenda."),
                    ("⚠️ Business Challenges & Pain Points", "business_challenges",
                     "What keeps the leadership awake at night? Regulatory, operational, competitive pressures."),
                    ("⚔️ Competitive Landscape", "competitive_landscape",
                     "Which other firms are engaged with this account? Where are we strong vs. at risk?"),
                ]
                for heading, field, placeholder in _si_sections:
                    val = _ap_c.get(field, "")
                    st.markdown(f"#### {heading}")
                    if val:
                        st.markdown(
                            f"<div style='background:#F5F8FA;border-left:3px solid #FF7A59;"
                            f"padding:10px 14px;border-radius:4px;font-size:13px;color:#33475B;"
                            f"white-space:pre-wrap'>{val}</div>",
                            unsafe_allow_html=True,
                        )
                    else:
                        st.caption(f"_{placeholder}_")
                    st.markdown("")

            # ── Stakeholders ──────────────────────────────────────────────────
            with _tab_sh:
                st.markdown("#### Stakeholder Map")

                _sh_groups = [
                    ("🏛️ Economic Buyer / Decision Maker", "#E8692A"),
                    ("🤝 Champion / Coach", "#00A862"),
                    ("🎖️ Executive Sponsor (Our Side)", "#0091AE"),
                    ("👥 Influencers & Users", "#516F90"),
                ]

                # Collect stakeholders from the account and its opportunities
                _dm_set = set()
                _inf_set = set()
                for o in _ap_opps:
                    if o.get("decision_maker"):
                        _dm_set.add(o["decision_maker"])
                    if o.get("influencers"):
                        _inf_set.add(o["influencers"])

                def _sh_card(label, value, color):
                    if value:
                        st.markdown(
                            f"<div style='background:white;border:1px solid #DFE3EB;"
                            f"border-left:3px solid {color};border-radius:6px;"
                            f"padding:10px 14px;margin-bottom:6px'>"
                            f"<div style='font-size:11px;color:#7C98B6;font-weight:600'>{label}</div>"
                            f"<div style='font-size:13px;color:#33475B;font-weight:600;margin-top:2px'>{value}</div>"
                            f"</div>",
                            unsafe_allow_html=True,
                        )

                _sh1, _sh2 = st.columns(2)
                with _sh1:
                    _sh_card("Key Contact", f"{_ap_c.get('key_contact','')} · {_ap_c.get('contact_title','')}" if _ap_c.get("key_contact") else "", "#0091AE")
                    for dm in _dm_set:
                        _sh_card("Decision Maker (from deals)", dm, "#E8692A")
                    _sh_card("Client Champion", _ap_c.get("champion", ""), "#00A862")
                with _sh2:
                    _sh_card("Our Executive Sponsor", _ap_c.get("executive_sponsor", ""), "#0091AE")
                    for inf in _inf_set:
                        _sh_card("Influencers (from deals)", inf, "#516F90")

                # EP/EM per deal
                _ep_em = {}
                for o in _ap_opps:
                    ep = o.get("engagement_partner", "")
                    em = o.get("engagement_manager", "")
                    if ep or em:
                        _ep_em[o["title"]] = (ep, em)
                if _ep_em:
                    st.markdown("#### Our Engagement Team")
                    for deal_title, (ep, em) in _ep_em.items():
                        st.markdown(
                            f"<div style='font-size:12px;color:#7C98B6;margin-bottom:2px'>{deal_title}</div>"
                            + (f"<span style='font-size:13px;color:#33475B;margin-right:16px'><b>EP:</b> {ep}</span>" if ep else "")
                            + (f"<span style='font-size:13px;color:#33475B'><b>EM:</b> {em}</span>" if em else ""),
                            unsafe_allow_html=True,
                        )
                        st.markdown("")

            # ── Deals & Projects ──────────────────────────────────────────────
            with _tab_deals:
                if _ap_active:
                    st.markdown(f"#### 🗂️ Active Opportunities ({len(_ap_active)})")
                    for o in sorted(_ap_active, key=lambda x: (
                        (x.get("engagement_partner") or "").lower() or "zzz",
                        x.get("expected_close_date") or "9999",
                    )):
                        color = STAGE_COLOR[o["stage"]]
                        prob_color = "#00A862" if o["probability"] >= 70 else "#F5A623" if o["probability"] >= 40 else "#E74C3C"
                        em_ep_d = " · ".join(filter(None, [
                            f"EP: {o['engagement_partner']}" if o.get("engagement_partner") else "",
                            f"EM: {o['engagement_manager']}" if o.get("engagement_manager") else "",
                        ]))
                        _dd1, _dd2 = st.columns([9, 1])
                        with _dd1:
                            st.markdown(
                                f"<div class='hs-deal' style='border-left:3px solid {color}'>"
                                f"<div style='display:flex;justify-content:space-between'>"
                                f"<span style='font-size:13px;font-weight:700;color:#33475B'>{o['title']}</span>"
                                f"<span style='font-size:15px;font-weight:700;color:#33475B;white-space:nowrap;margin-left:12px'>"
                                f"{_fmt_sgd(o['value_sgd'])}"
                                f"<span style='font-size:11px;color:{prob_color};margin-left:6px'>{o['probability']}%</span>"
                                f"</span></div>"
                                f"<div style='margin-top:4px;display:flex;gap:6px;align-items:center;flex-wrap:wrap'>"
                                + _stage_badge_html(o["stage"])
                                + (f"<span style='font-size:11px;color:#33475B'>{em_ep_d}</span>" if em_ep_d else "")
                                + (f"<span style='font-size:11px;color:#7C98B6'>📅 {o['expected_close_date']}</span>" if o.get("expected_close_date") else "")
                                + "</div>"
                                + (f"<div style='font-size:11px;color:#7C98B6;margin-top:3px'>⚡ {o['next_action']}</div>" if o.get("next_action") else "")
                                + "</div>",
                                unsafe_allow_html=True,
                            )
                        with _dd2:
                            if st.button("✏️", key=f"ap_ed_{o['id']}", use_container_width=True, help="Edit"):
                                st.session_state["open_opp_id"] = o["id"]
                                st.session_state.pop("edit_opp", None)
                                st.rerun()
                else:
                    st.info("No active opportunities for this account.")

                if _ap_won:
                    st.divider()
                    _won_color_ap = STAGE_COLOR["Won"]
                    st.markdown(f"#### ✅ Won Projects ({len(_ap_won)} · {_fmt_sgd(_ap_won_val)})")
                    for o in sorted(_ap_won, key=lambda x: (
                        (x.get("engagement_partner") or "").lower() or "zzz",
                        x.get("project_end_date") or x.get("expected_close_date") or "9999",
                    )):
                        em_ep_w = " · ".join(filter(None, [
                            f"EP: {o['engagement_partner']}" if o.get("engagement_partner") else "",
                            f"EM: {o['engagement_manager']}" if o.get("engagement_manager") else "",
                        ]))
                        st.markdown(
                            f"<div class='hs-won'>"
                            f"<div style='display:flex;justify-content:space-between'>"
                            f"<span style='font-weight:600;color:#33475B'>{o['title']}</span>"
                            f"<span style='font-size:15px;font-weight:700;color:{_won_color_ap}'>{_fmt_sgd(o['value_sgd'])}</span>"
                            f"</div>"
                            f"<div style='font-size:12px;color:#7C98B6;margin-top:2px'>"
                            + (f"{em_ep_w} &nbsp;·&nbsp; " if em_ep_w else "")
                            + (f"End {o['project_end_date']}" if o.get("project_end_date") else
                               (f"Closed {o['expected_close_date']}" if o.get("expected_close_date") else ""))
                            + "</div></div>",
                            unsafe_allow_html=True,
                        )

                if _ap_lost:
                    st.divider()
                    st.markdown(f"#### ❌ Lost Deals ({len(_ap_lost)})")
                    for o in _ap_lost:
                        st.markdown(
                            f"<div style='background:#F5F5F5;border-left:3px solid #9B9B9B;"
                            f"border-radius:4px;padding:8px 12px;margin-bottom:4px;"
                            f"font-size:13px;color:#888'>{o['title']} · {_fmt_sgd(o['value_sgd'])}</div>",
                            unsafe_allow_html=True,
                        )

            # ── Account Plan ──────────────────────────────────────────────────
            with _tab_ap:
                _ap_plan_sections = [
                    ("🎯 Our Account Goals", "account_goals",
                     "What are our 12-month revenue and relationship objectives for this account?"),
                    ("🚀 White Space & Growth Opportunities", "white_space",
                     "Which services have we not yet introduced? Where can we expand our footprint?"),
                ]
                for heading, field, placeholder in _ap_plan_sections:
                    val = _ap_c.get(field, "")
                    st.markdown(f"#### {heading}")
                    if val:
                        st.markdown(
                            f"<div style='background:#F5F8FA;border-left:3px solid #FF7A59;"
                            f"padding:10px 14px;border-radius:4px;font-size:13px;color:#33475B;"
                            f"white-space:pre-wrap'>{val}</div>",
                            unsafe_allow_html=True,
                        )
                    else:
                        st.caption(f"_{placeholder}_")
                    st.markdown("")

                # Upcoming actions from active deals
                _upcoming_acts = [
                    o for o in _ap_active if o.get("next_action")
                ]
                if _upcoming_acts:
                    st.markdown("#### ⚡ Next Steps")
                    for o in sorted(_upcoming_acts, key=lambda x: x.get("next_action_date") or "9999"):
                        d = _days_until(o.get("next_action_date", ""))
                        urgency = "🔴" if d is not None and d <= 3 else "🟡" if d is not None and d <= 7 else "🟢"
                        st.markdown(
                            f"{urgency} **{o['title']}** — {o['next_action']}"
                            + (f" · _{o['next_action_date']}_" if o.get("next_action_date") else "")
                        )

            # ── Edit Plan ─────────────────────────────────────────────────────
            with _tab_edit:
                with st.form("ap_edit_form"):
                    st.markdown("#### Account Profile")
                    _ep1, _ep2 = st.columns(2)
                    _ap_revenue = _ep1.text_input("Annual Revenue", value=_ap_c.get("annual_revenue", ""),
                                                  placeholder="e.g. S$500M")
                    _ap_employees = _ep2.text_input("No. of Employees", value=_ap_c.get("employee_count", ""),
                                                    placeholder="e.g. 5,000")
                    _ep3, _ep4 = st.columns(2)
                    _ap_website = _ep3.text_input("Website", value=_ap_c.get("website", ""),
                                                  placeholder="e.g. www.company.com")
                    _ap_budget_cycle = _ep4.text_input("Budget Cycle", value=_ap_c.get("budget_cycle", ""),
                                                       placeholder="e.g. April FY, S$2M AI budget")
                    _ep5, _ep6 = st.columns(2)
                    TIER_LIST = ["", "Strategic", "Key", "Growth", "Monitor"]
                    _ap_tier = _ep5.selectbox(
                        "Account Tier", TIER_LIST,
                        index=TIER_LIST.index(_ap_c.get("account_tier", "")) if _ap_c.get("account_tier", "") in TIER_LIST else 0,
                    )
                    _ap_rel = _ep6.slider("Relationship Score", 1, 5, _ap_c.get("relationship_score", 3))

                    st.markdown("#### Key Contacts")
                    _ec1, _ec2 = st.columns(2)
                    _ap_key_contact = _ec1.text_input("Key Contact", value=_ap_c.get("key_contact", ""))
                    _ap_contact_title = _ec2.text_input("Title / Role", value=_ap_c.get("contact_title", ""))
                    _ec3, _ec4 = st.columns(2)
                    _ap_exec_sponsor = _ec3.text_input("Our Executive Sponsor", value=_ap_c.get("executive_sponsor", ""),
                                                       placeholder="PwC partner sponsoring this account")
                    _ap_champion = _ec4.text_input("Client Champion", value=_ap_c.get("champion", ""),
                                                   placeholder="Client-side advocate for our work")

                    st.markdown("#### Strategic Intelligence")
                    _ap_strategy = st.text_area(
                        "Company Strategy & Priorities", value=_ap_c.get("business_strategy", ""),
                        height=100, placeholder="Key strategic initiatives, transformation agenda, growth priorities…",
                    )
                    _ap_challenges = st.text_area(
                        "Business Challenges & Pain Points", value=_ap_c.get("business_challenges", ""),
                        height=100, placeholder="Regulatory pressures, operational gaps, competitive threats…",
                    )
                    _ap_competitive = st.text_area(
                        "Competitive Landscape", value=_ap_c.get("competitive_landscape", ""),
                        height=80, placeholder="Other firms engaged, where we're strong vs. at risk…",
                    )

                    st.markdown("#### Account Plan")
                    _ap_goals = st.text_area(
                        "Our Account Goals (12-month)", value=_ap_c.get("account_goals", ""),
                        height=100, placeholder="Revenue target, relationship milestones, new service introductions…",
                    )
                    _ap_whitespace = st.text_area(
                        "White Space & Growth Opportunities", value=_ap_c.get("white_space", ""),
                        height=100, placeholder="Untapped service lines, new business units, upsell / cross-sell areas…",
                    )
                    _ap_notes = st.text_area("Notes", value=_ap_c.get("notes", ""), height=60)

                    _ef1, _ef2 = st.columns(2)
                    _ap_saved = _ef1.form_submit_button("💾 Save Account Plan", type="primary", use_container_width=True)
                    _ef2.form_submit_button("Cancel", use_container_width=True)

                    if _ap_saved:
                        db.upsert_client({
                            **_ap_c,
                            "annual_revenue": _ap_revenue,
                            "employee_count": _ap_employees,
                            "website": _ap_website,
                            "budget_cycle": _ap_budget_cycle,
                            "account_tier": _ap_tier,
                            "relationship_score": _ap_rel,
                            "key_contact": _ap_key_contact,
                            "contact_title": _ap_contact_title,
                            "executive_sponsor": _ap_exec_sponsor,
                            "champion": _ap_champion,
                            "business_strategy": _ap_strategy,
                            "business_challenges": _ap_challenges,
                            "competitive_landscape": _ap_competitive,
                            "account_goals": _ap_goals,
                            "white_space": _ap_whitespace,
                            "notes": _ap_notes,
                        })
                        st.success("Account plan saved.")
                        st.rerun()


# ─────────────────────────────────────────────────────────────────────────────
# CHANNEL PARTNERS
# ─────────────────────────────────────────────────────────────────────────────
elif page == "Channel Partners":
    CP_TYPES = [
        "Government Agency", "Industry Association", "Statutory Board",
        "Educational Institution", "Trade Association", "Professional Body", "Other",
    ]
    CP_STATUSES = ["Active", "Inactive", "MOU Signed", "In Negotiation", "Prospect", "Lapsed"]
    CP_TIERS = ["Strategic", "Gold", "Silver", "Standard"]

    st.markdown("# 🤝 Channel Partners")
    st.caption("Manage ecosystem partnerships — government agencies, trade bodies, and institutions.")

    _partners = db.get_channel_partners()

    # ── Summary metrics ────────────────────────────────────────────────────────
    _cp_active = [p for p in _partners if p["status"] in ("Active", "MOU Signed")]
    _cp_pipeline = sum(p.get("joint_pipeline_value") or 0 for p in _partners)
    _cp_referrals = sum(p.get("referrals_received") or 0 for p in _partners)
    _cp_converted = sum(p.get("referrals_converted") or 0 for p in _partners)
    _cp_conv_rate = round(_cp_converted / _cp_referrals * 100) if _cp_referrals else 0

    _cm1, _cm2, _cm3, _cm4 = st.columns(4)
    _cm1.metric("Total Partners", len(_partners))
    _cm2.metric("Active / MOU", len(_cp_active))
    _cm3.metric("Joint Pipeline", _fmt_sgd(_cp_pipeline))
    _cm4.metric("Referral Conversion", f"{_cp_conv_rate}%" if _cp_referrals else "—",
                help=f"{_cp_converted} converted from {_cp_referrals} referrals")

    # ── Add / Edit form ────────────────────────────────────────────────────────
    _cp_edit_key = "edit_cp"
    _cp1h, _cp1b = st.columns([5, 1])
    _cp1h.subheader("Partner Directory")
    if _cp1b.button("＋ Add Partner", type="primary", use_container_width=True, key="cp_new"):
        st.session_state[_cp_edit_key] = {}
        st.rerun()

    if _cp_edit_key in st.session_state:
        _cp_o = st.session_state[_cp_edit_key]
        _cp_form_title = "New Channel Partner" if not _cp_o.get("id") else f"Edit — {_cp_o.get('name', '')}"
        with st.expander(f"{'🆕' if not _cp_o.get('id') else '✏️'} {_cp_form_title}", expanded=True):
            with st.form("cp_form"):
                _cf1, _cf2 = st.columns(2)
                _cp_name = _cf1.text_input("Partner Name *", value=_cp_o.get("name", ""),
                                           placeholder="e.g. IMDA")
                _cp_type = _cf2.selectbox("Partner Type", CP_TYPES,
                                          index=CP_TYPES.index(_cp_o["partner_type"])
                                          if _cp_o.get("partner_type") in CP_TYPES else 0)

                _cf3, _cf4, _cf5 = st.columns(3)
                _cp_status = _cf3.selectbox("Status", CP_STATUSES,
                                            index=CP_STATUSES.index(_cp_o["status"])
                                            if _cp_o.get("status") in CP_STATUSES else 0)
                _cp_tier = _cf4.selectbox("Tier", CP_TIERS,
                                          index=CP_TIERS.index(_cp_o["tier"])
                                          if _cp_o.get("tier") in CP_TIERS else 3)
                _cp_pipeline_val = _cf5.number_input(
                    "Joint Pipeline (S$)", min_value=0.0,
                    value=float(_cp_o.get("joint_pipeline_value") or 0),
                    step=50_000.0, format="%.0f",
                )

                _cp_program = st.text_input("Program Name", value=_cp_o.get("program_name", ""),
                                            placeholder="e.g. IMDA Open Innovation Platform Partnership")

                _cf6, _cf7 = st.columns(2)
                _cp_ep = _cf6.text_input("EP (Engagement Partner)", value=_cp_o.get("engagement_partner", ""))
                _cp_em = _cf7.text_input("EM (Engagement Manager)", value=_cp_o.get("engagement_manager", ""))

                _cf8, _cf9, _cf10 = st.columns(3)
                _cp_contact = _cf8.text_input("Primary Contact", value=_cp_o.get("primary_contact", ""),
                                              placeholder="Name")
                _cp_ctitle = _cf9.text_input("Contact Title", value=_cp_o.get("contact_title", ""),
                                             placeholder="e.g. Director")
                _cp_cemail = _cf10.text_input("Contact Email", value=_cp_o.get("contact_email", ""))

                _cf11, _cf12 = st.columns(2)
                def _cp_date(val):
                    if val:
                        try:
                            return datetime.strptime(str(val)[:10], "%Y-%m-%d").date()
                        except ValueError:
                            pass
                    return None
                _cp_mou = _cf11.date_input("MOU / Agreement Date", value=_cp_date(_cp_o.get("mou_date")))
                _cp_renew = _cf12.date_input("Renewal Date", value=_cp_date(_cp_o.get("renewal_date")))

                _cf13, _cf14 = st.columns(2)
                _cp_ref_in = _cf13.number_input("Referrals Received", min_value=0,
                                                value=int(_cp_o.get("referrals_received") or 0))
                _cp_ref_conv = _cf14.number_input("Referrals Converted", min_value=0,
                                                  value=int(_cp_o.get("referrals_converted") or 0))

                _cp_sectors = st.text_input("Focus Sectors", value=_cp_o.get("focus_sectors", ""),
                                            placeholder="e.g. Financial Services, Government, Healthcare")
                _cp_notes = st.text_area("Notes", value=_cp_o.get("notes", ""), height=100,
                                         placeholder="Programme details, key initiatives, relationship context...")

                _cps1, _cps2 = st.columns(2)
                _cp_saved = _cps1.form_submit_button("💾 Save Partner", type="primary", use_container_width=True)
                _cp_cancel = _cps2.form_submit_button("Cancel", use_container_width=True)

                if _cp_saved and _cp_name:
                    db.upsert_channel_partner(dict(
                        id=_cp_o.get("id"),
                        name=_cp_name.strip(),
                        partner_type=_cp_type,
                        program_name=_cp_program,
                        status=_cp_status,
                        tier=_cp_tier,
                        engagement_partner=_cp_ep,
                        engagement_manager=_cp_em,
                        primary_contact=_cp_contact,
                        contact_title=_cp_ctitle,
                        contact_email=_cp_cemail,
                        mou_date=str(_cp_mou) if _cp_mou else "",
                        renewal_date=str(_cp_renew) if _cp_renew else "",
                        focus_sectors=_cp_sectors,
                        joint_pipeline_value=_cp_pipeline_val,
                        referrals_received=_cp_ref_in,
                        referrals_converted=_cp_ref_conv,
                        notes=_cp_notes,
                    ))
                    del st.session_state[_cp_edit_key]
                    st.rerun()
                if _cp_cancel:
                    del st.session_state[_cp_edit_key]
                    st.rerun()

    # ── Partner table (editable) ───────────────────────────────────────────────
    if not _partners:
        st.info("No channel partners yet — click '＋ Add Partner' to add your first one.")
    else:
        _cp_status_filter = st.selectbox("Filter by Status", ["All"] + CP_STATUSES, key="cp_status_filter",
                                         label_visibility="collapsed")
        _cp_display = _partners if _cp_status_filter == "All" else [
            p for p in _partners if p["status"] == _cp_status_filter
        ]

        if _cp_display:
            _cp_ids = [p["id"] for p in _cp_display]
            _cp_rows = []
            for p in _cp_display:
                _health_dot = {1: "🔴", 2: "🟠", 3: "🟡", 4: "🟢", 5: "🌟"}.get(
                    int(p.get("health_score") or 3), "🟡"
                )
                _status_dot = {"Active": "🟢", "MOU Signed": "🟢", "In Negotiation": "🟡",
                               "Prospect": "🔵", "Inactive": "⚫", "Lapsed": "🔴"}.get(
                    p.get("status", ""), "")
                _cp_rows.append({
                    "": _health_dot,
                    "Partner": p.get("name", ""),
                    "Type": p.get("partner_type") or "",
                    "Program": p.get("program_name") or "",
                    "Status": f"{_status_dot} {p.get('status', '')}".strip(),
                    "Tier": p.get("tier") or "",
                    "EP": p.get("engagement_partner") or "",
                    "Contact": p.get("primary_contact") or "",
                    "Next Meeting": (p.get("next_meeting_date") or "")[:10],
                    "Pipeline (S$)": float(p.get("joint_pipeline_value") or 0),
                })

            _cp_df = pd.DataFrame(_cp_rows)
            st.caption("Tap any row to open partner details and meeting prep.")
            _cp_sel_event = st.dataframe(
                _cp_df,
                use_container_width=True,
                hide_index=True,
                column_config={
                    "": st.column_config.TextColumn("", width="small"),
                    "Pipeline (S$)": st.column_config.NumberColumn("Pipeline (S$)", format="S$%,.0f"),
                },
                on_select="rerun",
                selection_mode="single-row",
                key="cp_tbl_sel",
            )
            _cp_sel_rows = _cp_sel_event.selection.rows
            _panel_cp_id = None
            if _cp_sel_rows:
                _panel_cp_id = _cp_ids[_cp_sel_rows[0]]
                st.session_state["open_cp_id"] = _panel_cp_id
                st.session_state.pop("edit_cp", None)
            elif st.session_state.get("open_cp_id") in _cp_ids:
                _panel_cp_id = st.session_state["open_cp_id"]
            if _panel_cp_id:
                st.divider()
                _cp_inline_panel(_panel_cp_id)

        # ── Renewals alert ─────────────────────────────────────────────────────
        _renewals_soon = [
            p for p in _partners
            if p.get("renewal_date") and 0 <= (_days_until(p["renewal_date"]) or 999) <= 90
        ]
        if _renewals_soon:
            st.divider()
            st.markdown("**⚠️ MOU / Agreement renewals due within 90 days:**")
            for p in sorted(_renewals_soon, key=lambda x: x.get("renewal_date", "")):
                d = _days_until(p["renewal_date"])
                badge = "🔴" if d <= 30 else "🟡"
                st.markdown(
                    f"{badge} **{p['name']}** — {p.get('program_name') or 'Partnership'} "
                    f"· renews {p['renewal_date']} ({d}d)"
                )


# ─────────────────────────────────────────────────────────────────────────────
# COMPANIES
# ─────────────────────────────────────────────────────────────────────────────
elif page == "Companies":
    _ch1, _ch2 = st.columns([6, 1])
    _ch1.markdown("# 🏢 Companies")
    if _ch2.button("＋ Add Company", type="primary", use_container_width=True):
        st.session_state["edit_client"] = {}
        st.session_state.pop("acra_prefill", None)

    _cf1, _cf2, _cf3 = st.columns([3, 2, 2])
    with _cf1:
        _search = st.text_input("Search", placeholder="Company name or contact...",
                                label_visibility="collapsed")
    with _cf2:
        _sectors = ["All Sectors"] + db.get_client_sectors()
        _sel_sector = st.selectbox("Sector", _sectors, label_visibility="collapsed")
    with _cf3:
        _sel_buyer = st.selectbox("Buyer Type", ["All Types", "Institutional", "Owner"],
                                  label_visibility="collapsed")

    # ── Add / Edit company form ────────────────────────────────────────────────
    if "edit_client" in st.session_state:
        c = st.session_state["edit_client"]
        _is_new = not c.get("id")

        with st.expander(
            f"{'🏢 New Company' if _is_new else '✏️ Edit — ' + c.get('company', '')}",
            expanded=True,
        ):
            if _is_new:
                st.markdown("**Step 1 — Search ACRA for the company**")
                _acra_col1, _acra_col2 = st.columns([4, 1])
                with _acra_col1:
                    _acra_query = st.text_input(
                        "Company name",
                        value=st.session_state.get("acra_query", ""),
                        placeholder="Type company name to search ACRA...",
                        key="acra_query_input",
                    )
                with _acra_col2:
                    st.write("")
                    _do_search = st.button("🔍 Search ACRA", use_container_width=True)

                if _do_search and _acra_query:
                    st.session_state["acra_query"] = _acra_query
                    with st.spinner("Searching ACRA..."):
                        results = _acra_search(_acra_query)
                    st.session_state["acra_results"] = results

                acra_results = st.session_state.get("acra_results", [])
                prefill = st.session_state.get("acra_prefill", {})

                if acra_results:
                    _labels = []
                    for r in acra_results:
                        _name = r.get("entity_name", r.get("name", r.get("company_name", "")))
                        _uen = r.get("uen", r.get("uen_no", ""))
                        _ssic = str(r.get("primary_ssic_code", r.get("ssic_code", r.get("ssic", "")))).strip()
                        _labels.append(f"{_name}"
                                       + (f" · UEN {_uen}" if _uen else "")
                                       + (f" · SSIC {_ssic}" if _ssic else ""))
                    _labels.insert(0, "(Enter manually — not in results)")
                    _selected = st.radio("Select matching company:", _labels, key="acra_pick")
                    if _selected != _labels[0]:
                        _idx = _labels.index(_selected) - 1
                        _r = acra_results[_idx]
                        _sel_name = _r.get("entity_name", _r.get("name", _r.get("company_name", "")))
                        _sel_ssic = str(_r.get("primary_ssic_code", _r.get("ssic_code", _r.get("ssic", "")))).strip()
                        _sel_sector_val = _ssic_to_sector(_sel_ssic) if _sel_ssic else "Other"
                        if st.button("✅ Use this company →", type="primary"):
                            st.session_state["acra_prefill"] = {
                                "company": _sel_name, "sector": _sel_sector_val,
                            }
                            st.session_state.pop("acra_results", None)
                            st.rerun()
                elif st.session_state.get("acra_query") and st.session_state.get("acra_results") is not None:
                    st.info("No ACRA results — enter details manually below.")

                if prefill:
                    st.success(
                        f"✅ ACRA match: **{prefill.get('company','')}** · {prefill.get('sector','')}"
                    )
                st.markdown("**Step 2 — Fill in company details**")
                st.divider()

            prefill = st.session_state.get("acra_prefill", {})
            with st.form("client_form"):
                company = st.text_input(
                    "Company Name *",
                    value=prefill.get("company", c.get("company", "")),
                )
                cc1, cc2 = st.columns(2)
                _default_sector = prefill.get("sector", c.get("sector", SECTOR_LIST[0]))
                sector = cc1.selectbox(
                    "Sector *", SECTOR_LIST,
                    index=SECTOR_LIST.index(_default_sector) if _default_sector in SECTOR_LIST else 0,
                )
                sub_sector = cc2.text_input("Sub-sector", value=c.get("sub_sector", ""))

                cc3, cc4 = st.columns(2)
                SIZE_LIST = ["Small (<200)", "Mid (1k-10k)", "Large (>10k)"]
                company_size = cc3.selectbox(
                    "Company Size", SIZE_LIST,
                    index=SIZE_LIST.index(c["company_size"]) if c.get("company_size") in SIZE_LIST else 2,
                )
                buyer_type = cc4.selectbox(
                    "Buyer Type *", ["Institutional", "Owner"],
                    index=0 if c.get("buyer_type", "Institutional") == "Institutional" else 1,
                )
                cc5, cc6 = st.columns(2)
                key_contact = cc5.text_input("Key Contact", value=c.get("key_contact", ""))
                contact_title = cc6.text_input("Title / Role", value=c.get("contact_title", ""))
                cc7, cc8 = st.columns(2)
                country = cc7.text_input("Country", value=c.get("country", "Singapore"))
                relationship_score = cc8.slider("Relationship Score", 1, 5, c.get("relationship_score", 3))
                MAT_LIST = ["Early", "Developing", "Intermediate", "Advanced"]
                ai_maturity = st.selectbox(
                    "AI Maturity", MAT_LIST,
                    index=MAT_LIST.index(c["ai_maturity"]) if c.get("ai_maturity") in MAT_LIST else 1,
                )
                notes = st.text_area("Notes", value=c.get("notes", ""), height=80)
                cs1, cs2 = st.columns(2)
                submitted = cs1.form_submit_button("💾 Save Company", type="primary", use_container_width=True)
                cancelled = cs2.form_submit_button("Cancel", use_container_width=True)

                if submitted and company:
                    db.upsert_client(dict(
                        id=c.get("id"), company=company, sector=sector,
                        sub_sector=sub_sector, company_size=company_size,
                        buyer_type=buyer_type, country=country,
                        key_contact=key_contact, contact_title=contact_title,
                        relationship_score=relationship_score, ai_maturity=ai_maturity,
                        notes=notes,
                    ))
                    del st.session_state["edit_client"]
                    st.session_state.pop("acra_prefill", None)
                    st.session_state.pop("acra_results", None)
                    st.session_state.pop("acra_query", None)
                    st.rerun()
                if cancelled:
                    del st.session_state["edit_client"]
                    st.session_state.pop("acra_prefill", None)
                    st.session_state.pop("acra_results", None)
                    st.session_state.pop("acra_query", None)
                    st.rerun()

    # ── Company cards ─────────────────────────────────────────────────────────
    clients = db.get_clients(
        sector=None if _sel_sector == "All Sectors" else _sel_sector,
        buyer_type=None if _sel_buyer == "All Types" else _sel_buyer,
    )
    if _search:
        _s = _search.lower()
        clients = [c for c in clients if _s in c["company"].lower()
                   or _s in (c.get("key_contact") or "").lower()]

    all_opps_cl = db.get_opportunities()
    all_acts_cl = db.get_engagements(limit=300)

    if not clients:
        st.info("No companies found — adjust filters or click '＋ Add Company' above.")

    for c in clients:
        opp_list = [o for o in all_opps_cl if o["client_id"] == c["id"]]
        active_c = [o for o in opp_list if o["stage"] not in ("Won", "Lost")]
        won_c = [o for o in opp_list if o["stage"] == "Won"]
        pipeline_val = sum(o["value_sgd"] or 0 for o in active_c)
        won_c_val = sum(o["value_sgd"] or 0 for o in won_c)
        buyer_badge = "🏛️ Institutional" if c["buyer_type"] == "Institutional" else "👤 Owner"
        rel_stars = "⭐" * c["relationship_score"]
        acts_c = [a for a in all_acts_cl if a["client_id"] == c["id"]]
        last_act = acts_c[0] if acts_c else None
        mat_color = {
            "Early": "#E74C3C", "Developing": "#F5A623",
            "Intermediate": "#0091AE", "Advanced": "#00A862",
        }.get(c.get("ai_maturity", ""), "#888")
        sector_color = STAGE_COLOR.get("Qualified", "#0091AE")

        with st.expander(
            f"**{c['company']}** · {c['sector']} · "
            f"{buyer_badge} · {rel_stars} · "
            f"Pipeline: **{_fmt_sgd(pipeline_val)}**"
            + (f" · Won: **{_fmt_sgd(won_c_val)}**" if won_c_val else "")
        ):
            _cl1, _cl2, _cl3 = st.columns([3, 2, 1])
            with _cl1:
                if c.get("key_contact"):
                    st.markdown(f"**👤 {c['key_contact']}** · {c.get('contact_title', '')}")
                if c.get("ai_maturity"):
                    st.markdown(
                        f"AI Maturity: <span style='color:{mat_color};font-weight:700'>"
                        f"{c['ai_maturity']}</span> &nbsp;·&nbsp; {c.get('company_size', '')} &nbsp;·&nbsp; {c.get('country', '')}",
                        unsafe_allow_html=True,
                    )
                if c.get("notes"):
                    st.caption(c["notes"])
            with _cl2:
                st.metric("Active Pipeline", _fmt_sgd(pipeline_val))
                if last_act:
                    st.caption(f"Last: {last_act['activity_type']} · {last_act['activity_date']}")
                if active_c:
                    st.caption(f"{len(active_c)} active deal(s) · " + ", ".join({o["stage"] for o in active_c}))
            with _cl3:
                if st.button("✏️ Edit", key=f"ed_c_{c['id']}", use_container_width=True):
                    st.session_state["edit_client"] = c
                    st.rerun()
                if st.button("＋ Deal", key=f"ad_c_{c['id']}", use_container_width=True,
                             type="primary"):
                    st.session_state["edit_opp"] = {"client_id": c["id"]}
                    st.rerun()
                if st.button("🗑️ Delete", key=f"dl_c_{c['id']}", use_container_width=True):
                    db.delete_client(c["id"])
                    st.rerun()

            if opp_list:
                st.markdown("**Deals**")
                _odf = pd.DataFrame(opp_list)[
                    ["title", "stage", "value_sgd", "probability", "expected_close_date", "next_action"]
                ]
                _odf.columns = ["Title", "Stage", "Value (S$)", "Prob %", "Close Date", "Next Action"]
                st.dataframe(_odf, use_container_width=True, hide_index=True)

            if acts_c:
                st.markdown("**Recent Activities**")
                for a in acts_c[:3]:
                    icon = ACTIVITY_ICONS.get(a["activity_type"], "📌")
                    st.markdown(
                        f"{icon} **{a['activity_type']}** · {a['activity_date']}"
                        + (f" — {a['summary'][:80]}" if a.get("summary") else "")
                    )

    # Handle ＋ Opportunity redirect from company card
    if "edit_opp" in st.session_state:
        with st.container(border=True):
            _deal_form(st.session_state["edit_opp"], db.get_clients(), form_key="opp_form_companies_new")

    if st.session_state.get("open_opp_id"):
        _opp_inline_panel(st.session_state["open_opp_id"])


# ─────────────────────────────────────────────────────────────────────────────
# ACTIVITIES
# ─────────────────────────────────────────────────────────────────────────────
elif page == "Activities":
    _ah1, _ah2 = st.columns([6, 1])
    _ah1.markdown("# 📅 Activities")
    if _ah2.button("＋ Log Activity", type="primary", use_container_width=True):
        st.session_state["act_form_open"] = True

    all_clients_act = db.get_clients()
    clients_map_act = {c["company"]: c["id"] for c in all_clients_act}

    with st.expander("➕ Log Activity", expanded=st.session_state.get("act_form_open", False)):
        with st.form("log_activity_form"):
            _la1, _la2 = st.columns(2)
            _co_sel = _la1.selectbox("Company *", list(clients_map_act.keys()) or ["(No companies yet)"])
            _act_type = _la2.selectbox("Activity Type *", ACTIVITY_TYPES)
            _cid_sel = clients_map_act.get(_co_sel)
            _opp_list_a = db.get_opportunities(client_id=_cid_sel) if _cid_sel else []
            _opp_map_a = {"(No linked deal)": None}
            _opp_map_a.update({o["title"]: o["id"] for o in _opp_list_a})
            _la3, _la4 = st.columns(2)
            _opp_sel = _la3.selectbox("Linked Deal", list(_opp_map_a.keys()))
            _act_date = _la4.date_input("Date *", value=date.today())
            _participants = st.text_input("Participants", placeholder="Names of people in the meeting/call")
            _summary = st.text_area("Summary / Notes *", height=80,
                                    placeholder="What happened? Key points discussed.")
            _la5, _la6 = st.columns(2)
            _outcomes = _la5.text_area("Outcomes / Decisions", height=60)
            _next_steps = _la6.text_area("Next Steps", height=60)
            _lfs1, _lfs2 = st.columns(2)
            _log_sub = _lfs1.form_submit_button("💾 Log Activity", type="primary", use_container_width=True)
            _log_can = _lfs2.form_submit_button("Cancel", use_container_width=True)
            if _log_sub and _summary and _cid_sel:
                db.add_engagement(dict(
                    client_id=_cid_sel, opportunity_id=_opp_map_a[_opp_sel],
                    activity_type=_act_type, activity_date=str(_act_date),
                    participants=_participants, summary=_summary,
                    outcomes=_outcomes, next_steps=_next_steps,
                ))
                st.session_state.pop("act_form_open", None)
                st.rerun()
            if _log_can:
                st.session_state.pop("act_form_open", None)
                st.rerun()

    st.divider()
    _af1, _af2 = st.columns(2)
    with _af1:
        _act_client_filter = st.selectbox(
            "Filter by Company", ["All Companies"] + list(clients_map_act.keys())
        )
    with _af2:
        _act_type_filter = st.selectbox("Activity Type", ["All Types"] + ACTIVITY_TYPES)

    _filter_cid = None if _act_client_filter == "All Companies" else clients_map_act.get(_act_client_filter)
    activities = db.get_engagements(client_id=_filter_cid, limit=100)
    if _act_type_filter != "All Types":
        activities = [a for a in activities if a["activity_type"] == _act_type_filter]

    if not activities:
        st.info("No activities logged yet — click '＋ Log Activity' to record your first one.")
    else:
        st.markdown(f"**{len(activities)} activities**")
        for a in activities:
            icon = ACTIVITY_ICONS.get(a["activity_type"], "📌")
            _aci, _acb = st.columns([1, 11])
            with _aci:
                st.markdown(f"### {icon}")
            with _acb:
                _hdr = f"**{a['activity_type']}** · {a['company']} · {a['activity_date']}"
                if a.get("opp_title"):
                    _hdr += f" · _{a['opp_title']}_"
                st.markdown(_hdr)
                if a.get("participants"):
                    st.caption(f"👥 {a['participants']}")
                if a.get("summary"):
                    st.markdown(a["summary"])
                if a.get("outcomes"):
                    st.markdown(f"**Outcome:** {a['outcomes']}")
                if a.get("next_steps"):
                    st.markdown(f"**Next:** {a['next_steps']}")
            st.divider()


# ─────────────────────────────────────────────────────────────────────────────
# EMINENCE
# ─────────────────────────────────────────────────────────────────────────────
elif page == "Eminence":
    st.markdown("# 🌟 Eminence")
    st.caption("Build your public profile as Singapore's #1 practical AI advisor")

    _em_tab_ai, _em_tab_plan, _em_tab_miles, _em_tab_tracker = st.tabs(
        ["🤖 AI Coach", "📅 Content Plan", "🎯 Milestones", "📊 Tracker"]
    )

    # ── AI Coach ──────────────────────────────────────────────────────────────
    with _em_tab_ai:
        api_key = os.environ.get("ANTHROPIC_API_KEY", "")
        if not api_key:
            st.warning("Set ANTHROPIC_API_KEY to use the AI coach.")
        else:
            _ec_clr, _ec_mdl = st.columns([1, 3])
            if _ec_clr.button("🗑️ Clear", key="clr_em_chat", help="Clear eminence chat history",
                               use_container_width=True):
                db.clear_eminence_chat_history()
                st.rerun()
            _em_chat_model = _CHAT_MODELS[_ec_mdl.selectbox(
                "Model", list(_CHAT_MODELS.keys()),
                index=list(_CHAT_MODELS.keys()).index(_CHAT_MODEL_DEFAULT),
                key="em_chat_model_sel", label_visibility="collapsed",
            )]

            _em_history = db.get_eminence_chat_history(limit=60)
            for _emm in _em_history:
                with st.chat_message(_emm["role"]):
                    st.markdown(_emm["content"])

            if not _em_history:
                st.info(
                    "👋 Ask me anything about building your eminence: draft LinkedIn posts, "
                    "plan content, prep op-ed angles, structure a speaking abstract, "
                    "or get your 100/200/300-day plan. I know your deals and deployments."
                )

            _em_prompt = st.chat_input("Ask your eminence coach…", key="em_chat_input")
            if _em_prompt:
                db.save_eminence_message("user", _em_prompt)
                with st.chat_message("user"):
                    st.markdown(_em_prompt)
                api_msgs = [{"role": m["role"], "content": m["content"]} for m in _em_history]
                api_msgs.append({"role": "user", "content": _em_prompt})
                with st.chat_message("assistant"):
                    with st.spinner("Thinking…"):
                        try:
                            _ec = anthropic.Anthropic(api_key=api_key)
                            _er = _ec.messages.create(
                                model=_em_chat_model,
                                max_tokens=4096,
                                system=_EMINENCE_SYSTEM,
                                messages=api_msgs,
                            )
                            _em_reply = _er.content[0].text
                        except Exception as _ex:
                            _em_reply = f"⚠️ Error: {_ex}"
                    st.markdown(_em_reply)
                db.save_eminence_message("assistant", _em_reply)
                st.rerun()

    # ── Content Plan ──────────────────────────────────────────────────────────
    with _em_tab_plan:
        _ARCHETYPES = ["Field Notes", "Pattern", "Board Implication", "Contrarian Bet"]
        _ARCHETYPE_ICONS = {
            "Field Notes": "📓", "Pattern": "🔍",
            "Board Implication": "🏛️", "Contrarian Bet": "🎲",
        }
        _ARCHETYPE_DESC = {
            "Field Notes": "Specific deployment story — show the mess, end with the lesson",
            "Pattern": "Synthesise across deployments — what you keep seeing that nobody talks about",
            "Board Implication": "CEO/board-level takeaway — what this means for capital allocation",
            "Contrarian Bet": "What you believed before the crowd — and what happened",
        }
        _CP_STATUSES = ["Draft", "In Progress", "Published", "Cancelled"]
        _CP_PLATFORMS = ["LinkedIn", "Business Times", "Straits Times", "Speaking", "Newsletter", "Podcast", "Other"]
        _CP_STATUS_COLOR = {
            "Draft": "#888", "In Progress": "#FF7A59", "Published": "#00A862", "Cancelled": "#ccc"
        }

        _cph1, _cph2 = st.columns([5, 1])
        _cph1.markdown("#### Content Queue")
        if _cph2.button("＋ Add", type="primary", key="add_cp_btn", use_container_width=True):
            st.session_state["add_content_plan"] = True

        if st.session_state.get("add_content_plan"):
            with st.expander("✍️ Plan a Content Piece", expanded=True):
                with st.form("content_plan_form"):
                    _cpa1, _cpa2 = st.columns(2)
                    _cp_arch = _cpa1.selectbox("Archetype *", _ARCHETYPES)
                    _cp_platform = _cpa2.selectbox("Platform", _CP_PLATFORMS)
                    _cp_title = st.text_input(
                        "Working Title *",
                        placeholder="e.g. What deploying an invoice agent at Nippon Paint taught me"
                    )
                    _cpb1, _cpb2 = st.columns(2)
                    _cp_date = _cpb1.date_input("Target Publish Date", value=None)
                    _cp_status = _cpb2.selectbox("Status", _CP_STATUSES)
                    _cp_notes = st.text_area(
                        "Notes / Key Points",
                        height=80,
                        placeholder="The real problem → what broke → the insight → what readers can steal"
                    )
                    _cps1, _cps2 = st.columns(2)
                    _cp_saved = _cps1.form_submit_button("💾 Save", type="primary", use_container_width=True)
                    _cp_cancel = _cps2.form_submit_button("Cancel", use_container_width=True)
                    if _cp_saved and _cp_title:
                        db.add_content_plan(dict(
                            archetype=_cp_arch, title=_cp_title,
                            platform=_cp_platform,
                            target_date=str(_cp_date) if _cp_date else "",
                            status=_cp_status, notes=_cp_notes,
                        ))
                        del st.session_state["add_content_plan"]
                        st.rerun()
                    if _cp_cancel:
                        del st.session_state["add_content_plan"]
                        st.rerun()

        _cp_status_filter = st.selectbox(
            "Filter by status", ["All"] + _CP_STATUSES, key="cp_status_filter",
            label_visibility="collapsed",
        )

        _cp_items = db.get_content_plan(status_filter=_cp_status_filter)

        if not _cp_items:
            st.info(
                "No content pieces planned yet. Use the AI Coach to draft ideas, "
                "then add them here to track from draft to published."
            )
        else:
            _by_arch = {}
            for _ci in _cp_items:
                _by_arch.setdefault(_ci["archetype"], []).append(_ci)
            for _arch in _ARCHETYPES:
                if _arch not in _by_arch:
                    continue
                _aicon = _ARCHETYPE_ICONS[_arch]
                st.markdown(f"**{_aicon} {_arch}** — _{_ARCHETYPE_DESC[_arch]}_")
                for _ci in _by_arch[_arch]:
                    _sc = _CP_STATUS_COLOR.get(_ci["status"], "#888")
                    with st.expander(f"{_ci['title']}", expanded=False):
                        _cr1, _cr2, _cr3 = st.columns([3, 2, 1])
                        _cr1.markdown(
                            f"<span style='background:{_sc}22;color:{_sc};padding:2px 8px;"
                            f"border-radius:10px;font-size:11px;font-weight:700'>"
                            f"{_ci['status'].upper()}</span>"
                            + (f" &nbsp; {_ci['platform']}" if _ci.get("platform") else ""),
                            unsafe_allow_html=True,
                        )
                        if _ci.get("target_date"):
                            _cr2.caption(f"📅 Target: {_ci['target_date']}")
                        with _cr3:
                            if st.button("🗑️", key=f"del_cp_{_ci['id']}", use_container_width=True):
                                db.delete_content_plan(_ci["id"])
                                st.rerun()
                        if _ci.get("notes"):
                            st.markdown(_ci["notes"])
                        _status_opts = [s for s in _CP_STATUSES if s != _ci["status"]]
                        _new_status = st.selectbox(
                            "Update status", ["— keep current —"] + _status_opts,
                            key=f"cp_status_{_ci['id']}", label_visibility="collapsed",
                        )
                        if _new_status != "— keep current —":
                            db.update_content_plan_status(_ci["id"], _new_status)
                            st.rerun()
                st.divider()

    # ── Milestones ────────────────────────────────────────────────────────────
    with _em_tab_miles:
        _MILESTONES = {
            "100-Day: Earn the Right to Speak": [
                ("mine_gold", "Write 'AI in a Box' framework doc with methodology + diagram"),
                ("client_stories", "Get 2 anonymised client outcome stories written (SETSCO, Everllence)"),
                ("declare_lane", "Publish 'My lane: Practical AI for Singapore's industrial & maritime economy'"),
                ("linkedin_cadence", "Ship first 4 LinkedIn essays (1/week): practitioner insight, not hot takes"),
                ("flagship_pov", "Write flagship POV paper: AI adoption barriers in industrial SMEs + ECI reality"),
                ("warm_ecosystem", "Reach out to SBF, NTU, SLA, Synapxe — offer to contribute, not headline"),
            ],
            "200-Day: Be Invited, Don't Ask": [
                ("speaking_sbf", "Land SBF speaking slot — deliver AI-in-a-Box framework live"),
                ("speaking_maritime", "Land maritime or manufacturing forum speaking slot"),
                ("benchmark_study", "Co-create NTU or SBF AI-adoption benchmark study — own the data"),
                ("bt_oped", "Publish BT/ST op-ed: 'Why Singapore's mid-market is the real AI battleground'"),
                ("maritime_beachhead", "Develop maritime-specific AI POV with IMC and Megastar"),
                ("content_20", "20+ pieces published on the 2:1:1 rotation (Field Notes : Pattern : Board)"),
            ],
            "300-Day: Own a Category": [
                ("annual_report", "Publish 'State of Practical AI in Singapore's Industrial Economy' report"),
                ("roundtable", "Host invite-only 15–20 leader roundtable — be the convener"),
                ("regional_stage", "Take one Singapore story to a regional Asia stage"),
                ("media_citations", "Achieve 3+ earned media citations as Singapore industrial AI authority"),
                ("referral_inbound", "Receive first inbound speaking/media invitation (not pitched by you)"),
                ("bench_development", "Clients quoting you in their own press/awards as their AI advisor"),
            ],
        }
        _PHASE_COLOR = {
            "100-Day: Earn the Right to Speak": "#0091AE",
            "200-Day: Be Invited, Don't Ask": "#FF7A59",
            "300-Day: Own a Category": "#00A862",
        }

        _em_checked = db.get_context("eminence_milestones") or ""
        _checked_set = set(_em_checked.split(",")) if _em_checked else set()

        st.markdown("#### 100 / 200 / 300-Day Plan")
        st.caption(
            "The compound eminence effect: grounded practitioner → be invited → own a category. "
            "Check off milestones as you complete them."
        )
        _any_changed = False
        _new_checked = set(_checked_set)

        for _phase, _items in _MILESTONES.items():
            _pc = _PHASE_COLOR[_phase]
            _done = sum(1 for k, _ in _items if k in _new_checked)
            st.markdown(
                f"<div style='background:{_pc}18;border-left:4px solid {_pc};"
                f"padding:8px 14px;border-radius:6px;margin-bottom:6px'>"
                f"<b style='color:{_pc}'>{_phase}</b> &nbsp;"
                f"<span style='font-size:12px;color:#888'>{_done}/{len(_items)} done</span></div>",
                unsafe_allow_html=True,
            )
            for _key, _label in _items:
                _val = st.checkbox(_label, value=_key in _new_checked, key=f"mile_{_key}")
                if _val and _key not in _new_checked:
                    _new_checked.add(_key)
                    _any_changed = True
                elif not _val and _key in _new_checked:
                    _new_checked.discard(_key)
                    _any_changed = True
            st.markdown("")

        if _any_changed:
            db.set_context("eminence_milestones", ",".join(sorted(_new_checked)))
            st.rerun()

        _total_done = sum(1 for _phase, _items in _MILESTONES.items()
                          for k, _ in _items if k in _new_checked)
        _total_all = sum(len(_items) for _, _items in _MILESTONES.items())
        if _total_done > 0:
            st.progress(_total_done / _total_all,
                        text=f"{_total_done}/{_total_all} milestones complete")

    # ── Tracker ───────────────────────────────────────────────────────────────
    with _em_tab_tracker:
        _eth1, _eth2 = st.columns([5, 1])
        _eth1.markdown("#### Published & Delivered")
        if _eth2.button("＋ Add Item", type="primary", key="add_em_item_btn", use_container_width=True):
            st.session_state["add_eminence"] = True

        em_type_filter = st.selectbox(
            "Filter by type",
            ["All", "Publication", "Speaking", "Event", "Award", "Media", "Advisory"],
            label_visibility="collapsed", key="em_type_filter",
        )

        if st.session_state.get("add_eminence"):
            with st.expander("✍️ Add Eminence Item", expanded=True):
                with st.form("em_form"):
                    em1, em2 = st.columns(2)
                    em_type = em1.selectbox(
                        "Type *", ["Publication", "Speaking", "Event", "Award", "Media", "Advisory"]
                    )
                    em_date = em2.date_input("Date *", value=date.today())
                    em_title = st.text_input("Title *", placeholder="Name of publication, talk, award...")
                    em3, em4 = st.columns(2)
                    em_sector = em3.selectbox("Sector", [
                        "Cross-sector", "Financial Services", "Government", "Technology",
                        "Healthcare", "Logistics & Supply Chain", "Real Estate",
                    ])
                    em_platform = em4.text_input("Platform / Venue", placeholder="e.g. Singapore FinTech Festival")
                    em_description = st.text_area("Description", height=80)
                    em5, em6 = st.columns(2)
                    em_impact = em5.slider("Impact Score", 1, 5, 3)
                    em_url = em6.text_input("URL (optional)")
                    es1, es2 = st.columns(2)
                    submitted = es1.form_submit_button("💾 Save", type="primary", use_container_width=True)
                    cancelled = es2.form_submit_button("Cancel", use_container_width=True)
                    if submitted and em_title:
                        db.add_eminence(dict(
                            type=em_type, title=em_title, date=str(em_date),
                            sector=em_sector, platform=em_platform,
                            description=em_description, impact_score=em_impact, url=em_url,
                        ))
                        del st.session_state["add_eminence"]
                        st.rerun()
                    if cancelled:
                        del st.session_state["add_eminence"]
                        st.rerun()

        eminence_items = db.get_eminence(type_filter=None if em_type_filter == "All" else em_type_filter)
        em_summary = db.get_eminence_summary()

        if em_summary:
            em_df = pd.DataFrame(em_summary)
            total_score = sum(e["impact_score"] for e in eminence_items)
            kcols = st.columns(len(em_df) + 1)
            kcols[0].metric("Total Impact Score", total_score)
            for i, row in em_df.iterrows():
                kcols[i + 1].metric(row["type"], int(row["count"]))
            st.divider()

        _EM_TYPE_ICONS = {
            "Publication": "📄", "Speaking": "🎤", "Event": "🎪",
            "Award": "🏆", "Media": "📺", "Advisory": "🏛️",
        }
        _EM_TYPE_COLOR = {
            "Publication": "#0091AE", "Speaking": "#FF7A59", "Event": "#F5A623",
            "Award": "#F5A623", "Media": "#516F90", "Advisory": "#00A862",
        }
        for em in eminence_items:
            icon = _EM_TYPE_ICONS.get(em["type"], "⭐")
            stars = "⭐" * em["impact_score"]
            tc = _EM_TYPE_COLOR.get(em["type"], "#888")
            with st.expander(f"{icon} **{em['title']}** · {em['date']} · {stars}"):
                er1, er2 = st.columns([4, 1])
                with er1:
                    st.markdown(
                        f"<span style='background:{tc}22;color:{tc};padding:2px 8px;"
                        f"border-radius:10px;font-size:11px;font-weight:700'>{em['type'].upper()}</span>"
                        + (f" &nbsp; {em['platform']}" if em.get("platform") else "")
                        + (f" &nbsp;·&nbsp; {em.get('sector','')}" if em.get("sector") else ""),
                        unsafe_allow_html=True,
                    )
                    if em.get("description"):
                        st.markdown(em["description"])
                    if em.get("url"):
                        st.markdown(f"[🔗 Link]({em['url']})")
                with er2:
                    if st.button("🗑️ Delete", key=f"del_em_{em['id']}", use_container_width=True):
                        db.delete_eminence(em["id"])
                        st.rerun()
